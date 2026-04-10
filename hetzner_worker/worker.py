"""
Hetzner bot worker — FastAPI service that runs on the Hetzner VPS.
Render API POSTs bot jobs here → worker runs linkedin bot → POSTs log lines back to Render.
Also handles tailor jobs: Render POSTs /run-tailor → Hetzner runs ResumeCrew → POSTs progress
back to Render callback URL.

Run: uvicorn worker:app --host 0.0.0.0 --port 8001
"""
from __future__ import annotations

import base64
import json
import os
import shutil
import subprocess
import sys
import tempfile
import threading
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

import httpx
from fastapi import BackgroundTasks, FastAPI, Header, HTTPException
from pydantic import BaseModel

app = FastAPI()

# Load .env file if present (so secret survives worker restarts without re-passing env var)
_env_file = Path(__file__).parent / ".env"
if _env_file.exists():
    for _line in _env_file.read_text().splitlines():
        _line = _line.strip()
        if _line and not _line.startswith("#") and "=" in _line:
            _k, _v = _line.split("=", 1)
            os.environ.setdefault(_k.strip(), _v.strip())

# Accept comma-separated list of secrets so multiple environments (Render, local) work simultaneously
_WORKER_SECRETS = {s.strip() for s in os.environ.get("WORKER_SECRET", "").split(",") if s.strip()}
WORKER_SECRET   = os.environ.get("WORKER_SECRET", "").split(",")[0].strip()  # kept for legacy use
_JOBEZEE_ROOT = Path(__file__).resolve().parent.parent
_BOT_DIR      = _JOBEZEE_ROOT / "linkedin_bot"
_LAUNCHER     = _BOT_DIR / "linkedin_launcher.py"

_procs: dict[str, subprocess.Popen] = {}
_lock  = threading.Lock()

# ── Concurrency limit + queue ─────────────────────────────────────────────────
MAX_CONCURRENT_BOTS = 1          # 1 bot at a time — 2GB RAM can't safely run 2 Chrome instances
_queue: list[BotJobRequest] = [] # jobs waiting for a free slot
_queue_lock = threading.Lock()

def _queue_position(job_id: str) -> int:
    """Return 1-based queue position, or 0 if not queued."""
    with _queue_lock:
        for i, j in enumerate(_queue):
            if j.job_id == job_id:
                return i + 1
    return 0

def _reap_dead_procs() -> int:
    """Remove processes from _procs that have already exited. Returns count removed."""
    reaped = 0
    with _lock:
        dead = [jid for jid, p in _procs.items() if p.poll() is not None]
        for jid in dead:
            del _procs[jid]
            reaped += 1
    return reaped

def _maybe_start_next() -> None:
    """If a slot is free, pop the next queued job and start it."""
    _reap_dead_procs()   # clean up zombies before checking slot count
    with _queue_lock:
        if not _queue:
            return
        with _lock:
            if len(_procs) >= MAX_CONCURRENT_BOTS:
                return
        job = _queue.pop(0)
    # Notify backend that job is starting (left the queue)
    try:
        _post_log(job.callback_url, job.job_id, "[JOBEZEE] Slot available — starting bot now")
    except Exception:
        pass
    t = threading.Thread(target=_run_bot_task, args=(job,), daemon=True)
    t.start()

# ── LinkedIn Connect sessions (per-user interactive login) ────────────────────
_connect_sessions: dict[str, dict] = {}   # user_id -> {proc, port}
_connect_lock = threading.Lock()
_CONNECT_CDP_PORT = 9223   # fixed CDP port for connect-session Chrome


# ── Auth ──────────────────────────────────────────────────────────────────────

def _auth(authorization: str) -> None:
    if _WORKER_SECRETS:
        token = authorization.removeprefix("Bearer ").strip()
        if token not in _WORKER_SECRETS:
            raise HTTPException(status_code=401, detail="Unauthorized")


# ── Models ────────────────────────────────────────────────────────────────────

class BotJobRequest(BaseModel):
    job_id:          str
    config:          dict
    resume_b64:      str = ""   # base64-encoded PDF
    resume_filename: str = ""
    callback_url:    str        # https://api.jobezee.org/api/bot/internal/log


# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    _reap_dead_procs()   # auto-clean before reporting counts
    with _lock:
        active = len(_procs)
    with _queue_lock:
        queued = len(_queue)
    with _connect_lock:
        connect_sessions = len(_connect_sessions)
    return {
        "status":           "ok",
        "active_jobs":      active,
        "queued_jobs":      queued,
        "max_concurrent":   MAX_CONCURRENT_BOTS,
        "connect_sessions": connect_sessions,   # separate from bot slots
    }


@app.post("/run-bot")
async def run_bot(
    job: BotJobRequest,
    background_tasks: BackgroundTasks,
    authorization: str = Header(...),
):
    _auth(authorization)
    _reap_dead_procs()   # don't queue behind zombies
    with _lock:
        active = len(_procs)
    if active < MAX_CONCURRENT_BOTS:
        background_tasks.add_task(_run_bot_task, job)
        return {"status": "started", "job_id": job.job_id}
    else:
        with _queue_lock:
            _queue.append(job)
            position = len(_queue)
        return {"status": "queued", "job_id": job.job_id, "queue_position": position}


@app.get("/queue-status/{job_id}")
def queue_status(job_id: str, authorization: str = Header(...)):
    """Check if a job is queued and its position."""
    _auth(authorization)
    pos = _queue_position(job_id)
    with _lock:
        active = len(_procs)
    return {
        "job_id":         job_id,
        "queued":         pos > 0,
        "queue_position": pos,
        "active_bots":    active,
        "max_bots":       MAX_CONCURRENT_BOTS,
    }


@app.get("/screenshot")
def screenshot(authorization: str = Header(...)):
    """Take a screenshot of the Xvfb display and return it as base64 PNG."""
    _auth(authorization)
    import subprocess, base64, tempfile
    tmp = tempfile.mktemp(suffix=".png")
    env = {**os.environ, "DISPLAY": ":99"}
    for cmd in [["scrot", tmp], ["import", "-window", "root", tmp]]:
        try:
            subprocess.run(cmd, env=env, timeout=5, check=True, capture_output=True)
            data = Path(tmp).read_bytes()
            # Extract PNG width/height without extra deps (bytes 16-24)
            width  = int.from_bytes(data[16:20], "big")
            height = int.from_bytes(data[20:24], "big")
            try:
                Path(tmp).unlink()
            except Exception:
                pass
            return {
                "image_b64": base64.b64encode(data).decode(),
                "format": "png",
                "width": width,
                "height": height,
            }
        except Exception:
            continue
    return {"image_b64": "", "error": "Screenshot tools not available (install scrot or imagemagick)"}


def _find_chrome_bin() -> str:
    import shutil as _sh
    env_bin = os.environ.get("CHROME_BIN", "")
    if env_bin and os.path.exists(env_bin):
        return env_bin
    for cb in (
        "/usr/bin/google-chrome",
        "/usr/bin/google-chrome-stable",
        "/opt/google/chrome/google-chrome",
        "/opt/google/chrome/chrome",
        "/usr/bin/chromium-browser",
        "/usr/bin/chromium",
    ):
        if os.path.exists(cb):
            return cb
    found = _sh.which("google-chrome") or _sh.which("google-chrome-stable") or _sh.which("chromium-browser") or _sh.which("chromium") or ""
    return found or "google-chrome"


def _connect_profile_dir(user_id: str) -> str:
    # Per-user isolated Chrome profile — prevents session cross-contamination
    import re as _re
    safe_id = _re.sub(r'[^a-zA-Z0-9-]', '', user_id)[:16] or "default"
    return os.path.expanduser(f"~/.config/google-chrome-linkedin-{safe_id}")


class ConnectStartRequest(BaseModel):
    user_id: str


@app.post("/connect/start")
def connect_start(req: ConnectStartRequest, authorization: str = Header(...)):
    """Launch a Chrome session on :99 for the user to log in to LinkedIn interactively."""
    _auth(authorization)

    # Kill any existing connect session for this user
    with _connect_lock:
        old = _connect_sessions.pop(req.user_id, None)
    if old:
        try:
            old["proc"].kill()
        except Exception:
            pass
    import time as _time, asyncio as _asyncio, json as _j
    from urllib.request import urlopen as _urlopen
    _time.sleep(0.5)

    profile_dir = _connect_profile_dir(req.user_id)
    os.makedirs(profile_dir, exist_ok=True)
    # Remove stale Chrome singleton locks
    for lock_file in ("SingletonLock", "SingletonSocket", "SingletonCookie"):
        lp = os.path.join(profile_dir, lock_file)
        try:
            if os.path.exists(lp) or os.path.islink(lp):
                os.remove(lp)
        except Exception:
            pass
    # Remove Chrome session files to prevent crash-restore bubble
    default_dir = os.path.join(profile_dir, "Default")
    for session_file in ("Last Session", "Last Tabs"):
        sf = os.path.join(default_dir, session_file)
        try:
            if os.path.exists(sf):
                os.remove(sf)
        except Exception:
            pass

    chrome_bin = _find_chrome_bin()
    cmd = [
        chrome_bin,
        f"--remote-debugging-port={_CONNECT_CDP_PORT}",
        "--no-sandbox",
        "--disable-dev-shm-usage",
        "--disable-gpu",
        "--disable-software-rasterizer",
        "--window-size=1280,800",
        "--no-first-run",
        "--no-default-browser-check",
        "--disable-extensions",
        "--disable-session-crashed-bubble",
        "--disable-features=ProfilePickerOnStartupFeature",
        "--disable-blink-features=AutomationControlled",
        f"--user-data-dir={profile_dir}",
        "--profile-directory=Default",
        "https://www.linkedin.com/login",
    ]
    env = {**os.environ, "DISPLAY": ":99"}
    proc = subprocess.Popen(cmd, env=env, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    with _connect_lock:
        _connect_sessions[req.user_id] = {"proc": proc, "port": _CONNECT_CDP_PORT}

    _time.sleep(4)   # give Chrome time to start + load page

    # Dismiss any LinkedIn "Something went wrong" popup
    async def _dismiss_popup():
        try:
            import websockets as _ws
            raw = _urlopen(f"http://127.0.0.1:{_CONNECT_CDP_PORT}/json", timeout=5).read()
            targets = _j.loads(raw)
            tab = next((t for t in targets if t.get("type") == "page"), None)
            if not tab:
                return
            async with _ws.connect(tab["webSocketDebuggerUrl"]) as sock:
                _id2 = [0]
                async def _s(method, params=None):
                    _id2[0] += 1
                    mid = _id2[0]
                    await sock.send(_j.dumps({"id": mid, "method": method, "params": params or {}}))
                    while True:
                        r = _j.loads(await _asyncio.wait_for(sock.recv(), 5))
                        if r.get("id") == mid:
                            return r
                await _s("Runtime.evaluate", {
                    "expression": """
                    (function() {
                        var btns = Array.from(document.querySelectorAll('button'));
                        for (var i = 0; i < btns.length; i++) {
                            var t = btns[i].innerText.trim().toLowerCase();
                            if (t === 'ok' || t === 'dismiss' || t === 'close' || t === 'got it') {
                                btns[i].click();
                            }
                        }
                        var dismiss = document.querySelector('[data-test-artdeco-toast-item-dismiss]')
                                   || document.querySelector('.artdeco-toast-item__dismiss')
                                   || document.querySelector('[aria-label="Dismiss"]');
                        if (dismiss) dismiss.click();
                        return 'done';
                    })()
                    """,
                    "returnByValue": True,
                })
        except Exception:
            pass
    _asyncio.run(_dismiss_popup())

    return {"status": "started", "port": _CONNECT_CDP_PORT}


class ConnectClickRequest(BaseModel):
    user_id: str
    x: int   # display pixel coords (1280x800 window)
    y: int


@app.post("/connect/click")
def connect_click(req: ConnectClickRequest, authorization: str = Header(...)):
    """Inject a mouse click at display coordinates via xdotool."""
    _auth(authorization)
    env = {**os.environ, "DISPLAY": ":99"}
    try:
        subprocess.run(
            ["xdotool", "mousemove", "--sync", str(req.x), str(req.y), "click", "1"],
            env=env, timeout=5, check=True, capture_output=True,
        )
    except Exception as e:
        raise HTTPException(500, f"xdotool click failed: {e}")
    return {"ok": True}


class ConnectTypeRequest(BaseModel):
    user_id: str
    text: str


@app.post("/connect/type")
def connect_type(req: ConnectTypeRequest, authorization: str = Header(...)):
    """Inject keyboard input via xdotool."""
    _auth(authorization)
    env = {**os.environ, "DISPLAY": ":99"}
    try:
        subprocess.run(
            ["xdotool", "type", "--clearmodifiers", "--delay", "50", req.text],
            env=env, timeout=15, check=True, capture_output=True,
        )
    except Exception as e:
        raise HTTPException(500, f"xdotool type failed: {e}")
    return {"ok": True}


class ConnectFillEmailRequest(BaseModel):
    user_id: str
    email: str


@app.post("/connect/fill-email")
async def connect_fill_email(req: ConnectFillEmailRequest, authorization: str = Header(...)):
    """Fill LinkedIn email field via CDP click + insertText so it renders visually."""
    _auth(authorization)
    import asyncio as _asyncio, json as _j
    from urllib.request import urlopen as _urlopen

    with _connect_lock:
        session = _connect_sessions.get(req.user_id)
    if not session:
        raise HTTPException(400, "No active connect session")

    try:
        import websockets as _ws
    except ImportError:
        raise HTTPException(500, "websockets not installed on worker")

    try:
        raw = _urlopen(f"http://127.0.0.1:{_CONNECT_CDP_PORT}/json", timeout=5).read()
        targets = _j.loads(raw)
        tab = next((t for t in targets if t.get("type") == "page"), None)
        if not tab:
            raise HTTPException(500, "No active Chrome tab found")

        _id = 0
        async def send(sock, method, params=None):
            nonlocal _id
            _id += 1
            msg_id = _id
            await sock.send(_j.dumps({"id": msg_id, "method": method, "params": params or {}}))
            while True:
                raw2 = _j.loads(await _asyncio.wait_for(sock.recv(), 8))
                if raw2.get("id") == msg_id:
                    return raw2

        async with _ws.connect(tab["webSocketDebuggerUrl"]) as sock:
            # 1) Wait up to 15s for the email field to appear (page may still be loading)
            coords = None
            for _ in range(30):
                res = await send(sock, "Runtime.evaluate", {
                    "expression": """
                    (function() {
                        var el = document.querySelector('#username')
                              || document.querySelector('input[name="session_key"]')
                              || document.querySelector('input[autocomplete="username"]')
                              || document.querySelector('input[type="email"]');
                        if (!el) return null;
                        var native = Object.getOwnPropertyDescriptor(window.HTMLInputElement.prototype, 'value');
                        native.set.call(el, '');
                        el.dispatchEvent(new Event('input',  {bubbles: true}));
                        el.dispatchEvent(new Event('change', {bubbles: true}));
                        var r = el.getBoundingClientRect();
                        return {x: r.left + r.width/2, y: r.top + r.height/2};
                    })()
                    """,
                    "returnByValue": True,
                })
                coords = (res.get("result", {}).get("result", {}) or {}).get("value")
                if coords:
                    break
                await _asyncio.sleep(0.5)
            if not coords:
                raise HTTPException(500, "Email field not found on page")

            x, y = coords["x"], coords["y"]

            # 2) Click to focus
            await send(sock, "Input.dispatchMouseEvent",
                       {"type": "mousePressed", "x": x, "y": y, "button": "left", "clickCount": 3})
            await _asyncio.sleep(0.1)
            await send(sock, "Input.dispatchMouseEvent",
                       {"type": "mouseReleased", "x": x, "y": y, "button": "left", "clickCount": 3})
            await _asyncio.sleep(0.15)

            # 3) Type — renders visually in the field and shows in screenshots
            await send(sock, "Input.insertText", {"text": req.email})

        return {"ok": True}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(500, str(exc))


class ConnectFillPasswordRequest(BaseModel):
    user_id: str
    password: str


@app.post("/connect/fill-password")
async def connect_fill_password(req: ConnectFillPasswordRequest, authorization: str = Header(...)):
    """Fill LinkedIn password field via CDP click + insertText so it renders visually."""
    _auth(authorization)
    import asyncio as _asyncio, json as _j
    from urllib.request import urlopen as _urlopen

    with _connect_lock:
        session = _connect_sessions.get(req.user_id)
    if not session:
        raise HTTPException(400, "No active connect session")

    try:
        import websockets as _ws
    except ImportError:
        raise HTTPException(500, "websockets not installed on worker")

    try:
        raw = _urlopen(f"http://127.0.0.1:{_CONNECT_CDP_PORT}/json", timeout=5).read()
        targets = _j.loads(raw)
        tab = next((t for t in targets if t.get("type") == "page"), None)
        if not tab:
            raise HTTPException(500, "No active Chrome tab found")

        _id = 0
        async def send(sock, method, params=None):
            nonlocal _id
            _id += 1
            msg_id = _id
            await sock.send(_j.dumps({"id": msg_id, "method": method, "params": params or {}}))
            while True:
                raw2 = _j.loads(await _asyncio.wait_for(sock.recv(), 8))
                if raw2.get("id") == msg_id:
                    return raw2

        async with _ws.connect(tab["webSocketDebuggerUrl"]) as sock:
            # 1) Wait up to 10s for password field to appear
            coords = None
            for _ in range(20):
                res = await send(sock, "Runtime.evaluate", {
                    "expression": """
                    (function() {
                        var el = document.querySelector('#password')
                              || document.querySelector('input[name="session_password"]')
                              || document.querySelector('input[type="password"]');
                        if (!el) return null;
                        var native = Object.getOwnPropertyDescriptor(window.HTMLInputElement.prototype, 'value');
                        native.set.call(el, '');
                        el.dispatchEvent(new Event('input',  {bubbles: true}));
                        el.dispatchEvent(new Event('change', {bubbles: true}));
                        var r = el.getBoundingClientRect();
                        return {x: r.left + r.width/2, y: r.top + r.height/2};
                    })()
                    """,
                    "returnByValue": True,
                })
                coords = (res.get("result", {}).get("result", {}) or {}).get("value")
                if coords:
                    break
                await _asyncio.sleep(0.5)
            if not coords:
                raise HTTPException(500, "Password field not found on page")

            x, y = coords["x"], coords["y"]
            # 2) Triple-click to focus + select all
            await send(sock, "Input.dispatchMouseEvent",
                       {"type": "mousePressed", "x": x, "y": y, "button": "left", "clickCount": 3})
            await _asyncio.sleep(0.1)
            await send(sock, "Input.dispatchMouseEvent",
                       {"type": "mouseReleased", "x": x, "y": y, "button": "left", "clickCount": 3})
            await _asyncio.sleep(0.15)
            # 3) Type password
            await send(sock, "Input.insertText", {"text": req.password})

        return {"ok": True}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(500, str(exc))


class ConnectFillCodeRequest(BaseModel):
    user_id: str
    code: str


@app.post("/connect/fill-code")
async def connect_fill_code(req: ConnectFillCodeRequest, authorization: str = Header(...)):
    """Fill verification/OTP code into any visible input and click Submit."""
    _auth(authorization)
    import asyncio as _asyncio, json as _j
    from urllib.request import urlopen as _urlopen

    with _connect_lock:
        session = _connect_sessions.get(req.user_id)
    if not session:
        raise HTTPException(400, "No active connect session")

    try:
        import websockets as _ws
    except ImportError:
        raise HTTPException(500, "websockets not installed on worker")

    try:
        raw = _urlopen(f"http://127.0.0.1:{_CONNECT_CDP_PORT}/json", timeout=5).read()
        targets = _j.loads(raw)
        tab = next((t for t in targets if t.get("type") == "page"), None)
        if not tab:
            raise HTTPException(500, "No active Chrome tab found")

        _id = 0
        async def send(sock, method, params=None):
            nonlocal _id
            _id += 1
            mid = _id
            await sock.send(_j.dumps({"id": mid, "method": method, "params": params or {}}))
            while True:
                raw2 = _j.loads(await _asyncio.wait_for(sock.recv(), 8))
                if raw2.get("id") == mid:
                    return raw2

        async with _ws.connect(tab["webSocketDebuggerUrl"]) as sock:
            # Find first visible text/number input, click it, clear it, type code
            res = await send(sock, "Runtime.evaluate", {
                "expression": """
                (function() {
                    var el = document.querySelector('input[type="text"]')
                          || document.querySelector('input[type="number"]')
                          || document.querySelector('input[type="tel"]')
                          || document.querySelector('input:not([type="hidden"]):not([type="submit"]):not([type="button"])');
                    if (!el) return null;
                    var native = Object.getOwnPropertyDescriptor(window.HTMLInputElement.prototype, 'value');
                    native.set.call(el, '');
                    el.dispatchEvent(new Event('input', {bubbles: true}));
                    var r = el.getBoundingClientRect();
                    return {x: r.left + r.width/2, y: r.top + r.height/2};
                })()
                """,
                "returnByValue": True,
            })
            coords = (res.get("result", {}).get("result", {}) or {}).get("value")
            if not coords:
                raise HTTPException(500, "Code input field not found on page")

            x, y = coords["x"], coords["y"]
            await send(sock, "Input.dispatchMouseEvent",
                       {"type": "mousePressed", "x": x, "y": y, "button": "left", "clickCount": 3})
            await _asyncio.sleep(0.1)
            await send(sock, "Input.dispatchMouseEvent",
                       {"type": "mouseReleased", "x": x, "y": y, "button": "left", "clickCount": 3})
            await _asyncio.sleep(0.15)
            await send(sock, "Input.insertText", {"text": req.code})
            await _asyncio.sleep(0.3)

            # Click Submit button
            res2 = await send(sock, "Runtime.evaluate", {
                "expression": """
                (function() {
                    var btn = document.querySelector('button[type="submit"]')
                           || Array.from(document.querySelectorAll('button')).find(function(b){
                               var t = b.innerText.trim().toLowerCase();
                               return t === 'submit' || t === 'verify' || t === 'continue';
                           });
                    if (!btn) return null;
                    var r = btn.getBoundingClientRect();
                    return {x: r.left + r.width/2, y: r.top + r.height/2};
                })()
                """,
                "returnByValue": True,
            })
            btn_coords = (res2.get("result", {}).get("result", {}) or {}).get("value")
            if btn_coords:
                bx, by = btn_coords["x"], btn_coords["y"]
                await send(sock, "Input.dispatchMouseEvent",
                           {"type": "mousePressed", "x": bx, "y": by, "button": "left", "clickCount": 1})
                await _asyncio.sleep(0.1)
                await send(sock, "Input.dispatchMouseEvent",
                           {"type": "mouseReleased", "x": bx, "y": by, "button": "left", "clickCount": 1})

        return {"ok": True}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(500, str(exc))


class ConnectPressLoginRequest(BaseModel):
    user_id: str


@app.post("/connect/press-login")
async def connect_press_login(req: ConnectPressLoginRequest, authorization: str = Header(...)):
    """Click Sign In button via CDP, poll URL, return result or captcha flag."""
    _auth(authorization)
    import asyncio as _asyncio, json as _j
    from urllib.request import urlopen as _urlopen

    with _connect_lock:
        session = _connect_sessions.get(req.user_id)
    if not session:
        raise HTTPException(400, "No active connect session")

    try:
        import websockets as _ws
    except ImportError:
        raise HTTPException(500, "websockets not installed on worker")

    _CAPTCHA_PATTERNS = ("checkpoint", "security", "challenge", "verification", "authwall")

    # Click Sign In button via CDP
    raw = _urlopen(f"http://127.0.0.1:{_CONNECT_CDP_PORT}/json", timeout=5).read()
    targets = _j.loads(raw)
    tab = next((t for t in targets if t.get("type") == "page"), None)
    if tab:
        async with _ws.connect(tab["webSocketDebuggerUrl"]) as sock:
            _id_c = 0
            async def _s(method, params=None):
                nonlocal _id_c
                _id_c += 1
                mid = _id_c
                await sock.send(_j.dumps({"id": mid, "method": method, "params": params or {}}))
                while True:
                    r = _j.loads(await _asyncio.wait_for(sock.recv(), 8))
                    if r.get("id") == mid:
                        return r
            res = await _s("Runtime.evaluate", {
                "expression": """
                (function() {
                    var btn = document.querySelector('button[data-litms-control-urn*="login"]')
                           || document.querySelector('button[aria-label*="Sign in"]')
                           || document.querySelector('.btn__primary--large')
                           || document.querySelector('button[type="submit"]');
                    if (!btn) return null;
                    var r = btn.getBoundingClientRect();
                    return {x: r.left + r.width/2, y: r.top + r.height/2};
                })()
                """,
                "returnByValue": True,
            })
            coords = (res.get("result", {}).get("result", {}) or {}).get("value")
            if coords:
                x, y = coords["x"], coords["y"]
                await _s("Input.dispatchMouseEvent",
                         {"type": "mousePressed", "x": x, "y": y, "button": "left", "clickCount": 1})
                await _asyncio.sleep(0.1)
                await _s("Input.dispatchMouseEvent",
                         {"type": "mouseReleased", "x": x, "y": y, "button": "left", "clickCount": 1})

    # Poll until URL leaves /login (max 20 s)
    async def _get_state():
        raw = _urlopen(f"http://127.0.0.1:{_CONNECT_CDP_PORT}/json", timeout=5).read()
        targets = _j.loads(raw)
        tab2 = next((t for t in targets if t.get("type") == "page"), None)
        if not tab2:
            return [], ""
        async with _ws.connect(tab2["webSocketDebuggerUrl"]) as sock2:
            _id_s = 0
            async def _s2(method, params=None):
                nonlocal _id_s
                _id_s += 1
                mid = _id_s
                await sock2.send(_j.dumps({"id": mid, "method": method, "params": params or {}}))
                while True:
                    r = _j.loads(await _asyncio.wait_for(sock2.recv(), 8))
                    if r.get("id") == mid:
                        return r
            r_url = await _s2("Runtime.evaluate",
                              {"expression": "window.location.href", "returnByValue": True})
            url = (r_url.get("result", {}).get("result", {}) or {}).get("value", "") or ""
            r_ck = await _s2("Network.getAllCookies")
            cookies = [c for c in r_ck.get("result", {}).get("cookies", [])
                       if "linkedin.com" in c.get("domain", "")]
        return cookies, url

    for _ in range(20):
        await _asyncio.sleep(1)
        try:
            _, interim_url = await _get_state()
            if "/login" not in interim_url:
                break
        except Exception:
            pass

    try:
        li_cookies, current_url = await _get_state()
    except Exception as exc:
        raise HTTPException(500, f"Cookie extraction failed: {exc}")

    has_session = any(c.get("name") == "li_at" for c in li_cookies)

    if not has_session and any(p in current_url for p in _CAPTCHA_PATTERNS):
        # Keep Chrome alive — user will solve CAPTCHA interactively
        return {"success": False, "captcha": True, "current_url": current_url,
                "message": "CAPTCHA detected — solve it in the screenshot panel"}

    # Success or definitive failure — kill Chrome
    with _connect_lock:
        s = _connect_sessions.pop(req.user_id, None)
    if s:
        try:
            s["proc"].kill()
        except Exception:
            pass

    return {
        "success": has_session,
        "cookies": li_cookies if has_session else [],
        "count": len(li_cookies) if has_session else 0,
        "current_url": current_url,
        "message": "Login successful" if has_session else "Login failed — wrong email or password",
    }


@app.get("/connect/cookies")
def connect_cookies(user_id: str, authorization: str = Header(...)):
    """Extract LinkedIn session cookies from the connect-session Chrome via CDP."""
    _auth(authorization)
    with _connect_lock:
        session = _connect_sessions.get(user_id)
    if not session:
        raise HTTPException(404, "No connect session for this user")

    port = session["port"]
    import asyncio as _asyncio
    import json as _j
    from urllib.request import urlopen as _urlopen

    async def _get_cookies_cdp():
        try:
            import websockets as _ws
        except ImportError:
            raise RuntimeError("websockets not installed on worker")

        raw = _urlopen(f"http://127.0.0.1:{port}/json", timeout=5).read()
        targets = _j.loads(raw)
        tab = next((t for t in targets if t.get("type") == "page"), None)
        if not tab:
            raise RuntimeError("No Chrome page found — is Chrome still running?")
        ws_url = tab["webSocketDebuggerUrl"]
        async with _ws.connect(ws_url) as sock:
            await sock.send(_j.dumps({"id": 1, "method": "Network.getAllCookies", "params": {}}))
            resp = await _asyncio.wait_for(sock.recv(), timeout=10)
            result = _j.loads(resp)
            return result.get("result", {}).get("cookies", [])

    try:
        all_cookies = _asyncio.run(_get_cookies_cdp())
    except Exception as e:
        raise HTTPException(500, f"Cookie extraction failed: {e}")

    li_cookies = [c for c in all_cookies if "linkedin.com" in c.get("domain", "")]
    return {"cookies": li_cookies, "count": len(li_cookies)}


@app.post("/connect/stop")
def connect_stop(user_id: str, authorization: str = Header(...)):
    """Kill the connect-session Chrome for this user."""
    _auth(authorization)
    with _connect_lock:
        session = _connect_sessions.pop(user_id, None)
    if session:
        try:
            session["proc"].kill()
        except Exception:
            pass
        return {"stopped": True}
    return {"stopped": False}


@app.get("/connect/status")
def connect_status(user_id: str, authorization: str = Header(...)):
    """Check if there's an active connect session for this user."""
    _auth(authorization)
    with _connect_lock:
        session = _connect_sessions.get(user_id)
    if not session:
        return {"active": False}
    proc = session["proc"]
    alive = proc.poll() is None
    if not alive:
        with _connect_lock:
            _connect_sessions.pop(user_id, None)
    return {"active": alive, "port": session.get("port")}


@app.get("/connect/page-url")
def connect_page_url(user_id: str, authorization: str = Header(...)):
    """Return the current page URL from the connect-session Chrome (used to detect CAPTCHA solved)."""
    _auth(authorization)
    import asyncio as _asyncio, json as _j
    from urllib.request import urlopen as _urlopen

    with _connect_lock:
        session = _connect_sessions.get(user_id)
    if not session:
        return {"url": ""}

    port = session.get("port", _CONNECT_CDP_PORT)

    async def _get_url():
        try:
            import websockets as _ws
        except ImportError:
            return ""
        try:
            raw = _urlopen(f"http://127.0.0.1:{port}/json", timeout=3).read()
            tab = next((t for t in _j.loads(raw) if t.get("type") == "page"), None)
            if not tab:
                return ""
            async with _ws.connect(tab["webSocketDebuggerUrl"]) as sock:
                await sock.send(_j.dumps({"id": 1, "method": "Runtime.evaluate",
                                           "params": {"expression": "window.location.href",
                                                      "returnByValue": True}}))
                r = _j.loads(await _asyncio.wait_for(sock.recv(), 5))
                return r.get("result", {}).get("result", {}).get("value", "")
        except Exception:
            return ""

    try:
        url = _asyncio.run(_get_url())
    except Exception:
        url = ""
    return {"url": url}


class ConnectDoLoginRequest(BaseModel):
    user_id: str
    email: str
    password: str


@app.post("/connect/do-login")
def connect_do_login(req: ConnectDoLoginRequest, authorization: str = Header(...)):
    """
    All-in-one: start Chrome on bot profile, navigate to LinkedIn login,
    fill email+password via CDP JS injection, wait for redirect, return cookies.
    Uses the bot's persistent profile so cookies are stored natively.
    """
    _auth(authorization)
    import time as _t, asyncio as _asyncio, json as _j
    from urllib.request import urlopen as _urlopen

    # Kill any existing connect session
    with _connect_lock:
        old = _connect_sessions.pop(req.user_id, None)
    if old:
        try:
            old["proc"].kill()
        except Exception:
            pass
    _t.sleep(0.5)

    profile_dir = _connect_profile_dir(req.user_id)
    os.makedirs(profile_dir, exist_ok=True)
    for lf in ("SingletonLock", "SingletonSocket", "SingletonCookie"):
        lp = os.path.join(profile_dir, lf)
        try:
            if os.path.exists(lp) or os.path.islink(lp):
                os.remove(lp)
        except Exception:
            pass

    chrome_bin = _find_chrome_bin()
    cmd = [
        chrome_bin,
        f"--remote-debugging-port={_CONNECT_CDP_PORT}",
        "--no-sandbox", "--disable-dev-shm-usage", "--disable-gpu",
        "--window-size=1280,800", "--no-first-run", "--no-default-browser-check",
        "--disable-extensions", "--disable-blink-features=AutomationControlled",
        "--disable-features=ProfilePickerOnStartupFeature",
        "--no-restore-last-session",
        f"--user-data-dir={profile_dir}", "--profile-directory=Default",
        "https://www.linkedin.com/login",
    ]
    env = {**os.environ, "DISPLAY": ":99"}
    proc = subprocess.Popen(cmd, env=env, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    with _connect_lock:
        _connect_sessions[req.user_id] = {"proc": proc, "port": _CONNECT_CDP_PORT}

    _t.sleep(5)  # wait for Chrome + LinkedIn page to load

    # Dismiss "Restore pages?" dialog with OS-level Escape key
    subprocess.run(["xdotool", "key", "--clearmodifiers", "Escape"], env=env,
                   capture_output=True)
    _t.sleep(1)

    def _xdotool_login():
        """Fill LinkedIn login form using xdotool — indistinguishable from real input."""
        r = subprocess.run(
            ["xdotool", "search", "--onlyvisible", "--class", "google-chrome"],
            env=env, capture_output=True, text=True
        )
        win_ids = [w for w in r.stdout.strip().split() if w]
        if not win_ids:
            raise RuntimeError("Chrome window not found — xdotool search returned nothing")
        win_id = win_ids[-1]

        subprocess.run(["xdotool", "windowactivate", "--sync", win_id], env=env,
                       capture_output=True)
        _t.sleep(0.4)

        geo_r = subprocess.run(
            ["xdotool", "getwindowgeometry", "--shell", win_id],
            env=env, capture_output=True, text=True
        )
        geo = {}
        for line in geo_r.stdout.strip().splitlines():
            if '=' in line:
                k, v = line.split('=', 1)
                try:
                    geo[k.strip()] = int(v.strip())
                except ValueError:
                    pass
        win_x = geo.get('X', 0)
        win_y = geo.get('Y', 0)
        win_w = geo.get('WIDTH', 1280)
        # Chrome title-bar(35) + address-bar(50) + --no-sandbox info-bar(30) = 115px
        header = 115

        # LinkedIn login page: email field center at viewport y≈415 (confirmed via screenshot)
        cx      = win_x + win_w // 2
        email_y = win_y + header + 415

        # Click email field, clear any pre-filled text, type email
        subprocess.run(["xdotool", "mousemove", str(cx), str(email_y)], env=env,
                       capture_output=True)
        subprocess.run(["xdotool", "click", "1"], env=env, capture_output=True)
        _t.sleep(0.3)
        subprocess.run(["xdotool", "key", "ctrl+a"], env=env, capture_output=True)
        _t.sleep(0.1)
        subprocess.run(["xdotool", "type", "--clearmodifiers", "--delay", "40",
                         req.email], env=env, capture_output=True)
        _t.sleep(0.5)

        # Tab to password field (more reliable than coordinate click)
        subprocess.run(["xdotool", "key", "Tab"], env=env, capture_output=True)
        _t.sleep(0.3)
        subprocess.run(["xdotool", "type", "--clearmodifiers", "--delay", "40",
                         req.password], env=env, capture_output=True)
        _t.sleep(0.3)

        # Press Enter to submit the form
        subprocess.run(["xdotool", "key", "Return"], env=env, capture_output=True)

    async def _read_result_after_login():
        """CDP used only for reading URL + cookies — no JS page interaction."""
        try:
            import websockets as _ws
        except ImportError:
            raise RuntimeError("websockets package not installed on worker")
        raw = _urlopen(f"http://127.0.0.1:{_CONNECT_CDP_PORT}/json", timeout=5).read()
        targets = _j.loads(raw)
        tab = next((t for t in targets if t.get("type") == "page"), None)
        if not tab:
            raise RuntimeError("No Chrome page found")
        async with _ws.connect(tab["webSocketDebuggerUrl"]) as sock:
            await sock.send(_j.dumps({"id": 10, "method": "Runtime.evaluate",
                                       "params": {"expression": "window.location.href",
                                                  "returnByValue": True}}))
            r_url = _j.loads(await _asyncio.wait_for(sock.recv(), timeout=8))
            current_url = (r_url.get("result", {}).get("result", {}).get("value", "") or "")
            await sock.send(_j.dumps({"id": 11, "method": "Network.getAllCookies",
                                       "params": {}}))
            r_ck = _j.loads(await _asyncio.wait_for(sock.recv(), timeout=10))
            all_cookies = r_ck.get("result", {}).get("cookies", [])
        li_cookies = [c for c in all_cookies if "linkedin.com" in c.get("domain", "")]
        return li_cookies, current_url

    def _kill_connect_chrome():
        with _connect_lock:
            s = _connect_sessions.pop(req.user_id, None)
        if s:
            try:
                s["proc"].kill()
            except Exception:
                pass

    try:
        _xdotool_login()
    except Exception as exc:
        _kill_connect_chrome()
        raise HTTPException(500, f"Login automation failed: {exc}")

    # Poll URL until it leaves /login (max 20 s) — more reliable than fixed sleep
    _CAPTCHA_PATTERNS = ("checkpoint", "security", "challenge", "verification", "authwall")
    for _ in range(20):
        _t.sleep(1)
        try:
            _interim = _asyncio.run(_read_result_after_login())
            _interim_url = _interim[1]
            if "/login" not in _interim_url:
                break
        except Exception:
            pass

    try:
        li_cookies, current_url = _asyncio.run(_read_result_after_login())
    except Exception as exc:
        _kill_connect_chrome()
        raise HTTPException(500, f"Cookie extraction failed: {exc}")

    has_session = any(c.get("name") == "li_at" for c in li_cookies)

    # CAPTCHA / security challenge — keep Chrome alive so the user can solve it interactively
    if not has_session and any(p in current_url for p in _CAPTCHA_PATTERNS):
        return {
            "success": False,
            "captcha": True,
            "current_url": current_url,
            "message": "CAPTCHA detected — solve it in the screenshot panel",
        }

    # Success or definitive failure — kill Chrome now
    _kill_connect_chrome()

    return {
        "success": has_session,
        "cookies": li_cookies if has_session else [],
        "count": len(li_cookies) if has_session else 0,
        "current_url": current_url,
        "message": "Login successful" if has_session else "Login failed — check your email and password",
    }


@app.post("/git-pull")
def git_pull(authorization: str = Header(...)):
    """Pull latest code from origin and return output. Force-resets to origin/main."""
    _auth(authorization)
    try:
        fetch = subprocess.run(
            ["git", "fetch", "origin"],
            cwd=str(_JOBEZEE_ROOT),
            capture_output=True, text=True, timeout=60,
        )
        reset = subprocess.run(
            ["git", "reset", "--hard", "origin/main"],
            cwd=str(_JOBEZEE_ROOT),
            capture_output=True, text=True, timeout=60,
        )
        return {
            "returncode": reset.returncode,
            "stdout": (fetch.stdout + reset.stdout).strip(),
            "stderr": (fetch.stderr + reset.stderr).strip(),
        }
    except Exception as e:
        return {"returncode": -1, "stdout": "", "stderr": str(e)}


def _kill_all_bot_procs() -> int:
    """Kill every tracked bot process + any orphaned Chrome/python-bot processes.
    Returns number of proc entries that were killed."""
    import signal as _sig
    killed = 0
    with _lock:
        job_ids = list(_procs.keys())
    for jid in job_ids:
        with _lock:
            proc = _procs.pop(jid, None)
        if proc is None:
            continue
        try:
            os.killpg(os.getpgid(proc.pid), _sig.SIGKILL)
            killed += 1
        except Exception:
            try:
                proc.kill()
                killed += 1
            except Exception:
                pass
    # Kill any orphaned Chrome sessions that the bot may have opened
    for pattern in ["linkedin_launcher", "runAiBot", "google-chrome.*--user-data-dir"]:
        try:
            subprocess.run(["pkill", "-9", "-f", pattern], capture_output=True, timeout=5)
        except Exception:
            pass
    # Also drain the queue so nothing restarts after a stop-all
    with _queue_lock:
        _queue.clear()
    return killed


@app.post("/clear-zombies")
def clear_zombies(authorization: str = Header(...)):
    """Remove dead processes from the active job table and start any queued jobs."""
    _auth(authorization)
    reaped = _reap_dead_procs()
    _maybe_start_next()
    with _lock:
        active = len(_procs)
    with _queue_lock:
        queued = len(_queue)
    return {"reaped": reaped, "active_jobs": active, "queued_jobs": queued}


@app.post("/stop-bot/{job_id}")
def stop_bot(job_id: str, authorization: str = Header(...)):
    _auth(authorization)
    import signal as _sig
    with _lock:
        proc = _procs.pop(job_id, None)
    if proc:
        try:
            os.killpg(os.getpgid(proc.pid), _sig.SIGKILL)
        except Exception:
            try:
                proc.kill()
            except Exception:
                pass
        # Kill orphaned Chrome for good measure
        for pattern in ["linkedin_launcher", "runAiBot"]:
            try:
                subprocess.run(["pkill", "-9", "-f", pattern], capture_output=True, timeout=5)
            except Exception:
                pass
        return {"stopped": True}
    # Job not in _procs — might have already finished; still do a broad kill so stray
    # Chrome processes don't keep running
    _kill_all_bot_procs()
    return {"stopped": False, "fallback_kill": True}


@app.post("/stop-all")
def stop_all(authorization: str = Header(...)):
    """Kill every running bot (used as fallback when job_id is lost, e.g. after Render restart)."""
    _auth(authorization)
    killed = _kill_all_bot_procs()
    return {"stopped": True, "killed": killed}


# ── Bot runner ────────────────────────────────────────────────────────────────

def _post_log(callback_url: str, job_id: str, line: str, status: str = "running") -> None:
    """POST a log line to the Render callback URL.

    httpx drops Authorization on cross-domain redirects (RFC 9110 security).
    We manually follow one redirect level to preserve the header.
    """
    headers = {"Authorization": f"Bearer {WORKER_SECRET}"}
    body    = {"job_id": job_id, "line": line, "status": status}
    try:
        resp = httpx.post(callback_url, json=body, headers=headers, timeout=5, follow_redirects=False)
        if resp.is_redirect:
            loc = resp.headers.get("location", "")
            if loc:
                resp2 = httpx.post(loc, json=body, headers=headers, timeout=5, follow_redirects=False)
                if resp2.status_code not in (200, 201, 204):
                    print(f"[_post_log] callback {loc} returned {resp2.status_code}: {resp2.text[:120]}", flush=True)
        elif resp.status_code not in (200, 201, 204):
            print(f"[_post_log] callback {callback_url} returned {resp.status_code}: {resp.text[:120]}", flush=True)
    except Exception as e:
        print(f"[_post_log] failed to post log to {callback_url}: {e}", flush=True)


def _run_bot_task(job: BotJobRequest) -> None:
    tmp_config     = None
    tmp_resume_dir = None

    try:
        # Write resume PDF to temp file
        resume_path = ""
        if job.resume_b64 and job.resume_filename:
            tmp_resume_dir = tempfile.mkdtemp(prefix="jobezee_resume_")
            resume_path    = str(Path(tmp_resume_dir) / job.resume_filename)
            with open(resume_path, "wb") as f:
                f.write(base64.b64decode(job.resume_b64))

        # Patch config for Linux/Hetzner
        config = job.config
        if resume_path and "questions" in config:
            config["questions"]["default_resume_path"] = resume_path
        if "settings" in config:
            config["settings"]["run_in_background"] = False  # use Xvfb display, not headless
            config["settings"]["stealth_mode"]      = False  # regular selenium — uc incompatible with Chrome 146
            config["settings"]["safe_mode"]         = False  # use persistent bot profile
            config["settings"]["jobezee_root"]      = str(_JOBEZEE_ROOT)

        # Write config JSON
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".json", delete=False, encoding="utf-8"
        ) as f:
            json.dump(config, f)
            tmp_config = f.name

        _uid         = config.get("settings", {}).get("user_profile_id", "")
        _openai_key  = config.get("settings", {}).get("openai_api_key",  "") or os.environ.get("OPENAI_API_KEY",    "")
        _claude_key  = config.get("settings", {}).get("claude_api_key",  "") or os.environ.get("CLAUDE_API_KEY",    "") or os.environ.get("ANTHROPIC_API_KEY", "")
        env = {
            **os.environ,
            "PYTHONUNBUFFERED": "1",
            "PYTHONIOENCODING": "utf-8",
            "PYTHONUTF8":       "1",
            "DISPLAY":          ":99",   # Xvfb virtual display
            "TWOCAPTCHA_API_KEY": os.environ.get("TWOCAPTCHA_API_KEY", ""),
            "JOBEZEE_USER_ID":  _uid,          # per-user Chrome profile isolation
            "OPENAI_API_KEY":   _openai_key,   # tools 1, 3, 4 (keyword, evaluator, latex)
            "CLAUDE_API_KEY":   _claude_key,   # tool 2 (ResumeTailor — Anthropic)
            "ANTHROPIC_API_KEY": _claude_key,  # fallback alias used by some anthropic SDK versions
        }

        cmd  = [sys.executable, "-u", str(_LAUNCHER), "--config", tmp_config]
        proc = subprocess.Popen(
            cmd,
            cwd=str(_BOT_DIR),
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1,
            encoding="utf-8",
            errors="replace",
            env=env,
            start_new_session=True,  # new process group — lets killpg kill launcher + chrome
        )

        with _lock:
            _procs[job.job_id] = proc

        # ── Heartbeat thread: posts a ping every 30 s so Render knows bot is alive.
        # If the worker crashes (OOM etc.) the heartbeat stops and Render can detect it.
        import time as _time
        _hb_stop = threading.Event()
        def _heartbeat():
            while not _hb_stop.wait(30):
                if proc.poll() is not None:
                    break   # process already finished
                _post_log(job.callback_url, job.job_id, "[JOBEZEE] [heartbeat]")
        _hb_thread = threading.Thread(target=_heartbeat, daemon=True)
        _hb_thread.start()

        for line in proc.stdout:
            line = line.rstrip()
            if line:
                _post_log(job.callback_url, job.job_id, line)

        _hb_stop.set()
        proc.wait()
        exit_code = proc.returncode
        status    = "complete" if exit_code == 0 else "error"
        _post_log(job.callback_url, job.job_id, f"[JOBEZEE] Bot finished (exit code {exit_code})", status)

    except Exception as exc:
        _post_log(job.callback_url, job.job_id, f"[JOBEZEE] Worker ERROR: {exc}", "error")
    finally:
        with _lock:
            _procs.pop(job.job_id, None)
        if tmp_config:
            try:
                Path(tmp_config).unlink()
            except Exception:
                pass
        if tmp_resume_dir:
            try:
                import shutil
                shutil.rmtree(tmp_resume_dir, ignore_errors=True)
            except Exception:
                pass
        # Slot freed — start next queued job if any
        _maybe_start_next()


# ══════════════════════════════════════════════════════════════════════════════
# TAILOR JOBS — ResumeCrew runs here, progress POSTed back to Render callback
# ══════════════════════════════════════════════════════════════════════════════

def _tailor_max_workers() -> int:
    try:
        import psutil
        available_gb = psutil.virtual_memory().available / (1024 ** 3)
        return max(1, min(6, int(available_gb / 0.5)))
    except Exception:
        return 4

_tailor_executor: ThreadPoolExecutor | None = None
_tailor_executor_lock = threading.Lock()

def _get_tailor_executor() -> ThreadPoolExecutor:
    global _tailor_executor
    with _tailor_executor_lock:
        if _tailor_executor is None or _tailor_executor._shutdown:
            n = _tailor_max_workers()
            _tailor_executor = ThreadPoolExecutor(max_workers=n, thread_name_prefix="tailor")
            print(f"[tailor_worker] Thread pool: {n} slots")
    return _tailor_executor


class TailorJobRequest(BaseModel):
    job_id:        str
    job_description: str
    resume_text:   str
    openai_api_key: str = ""
    username:      str = "resume"
    company:       str = "company"
    callback_url:  str   # https://api.jobezee.org/api/tailor/internal/callback


def _post_tailor_callback(callback_url: str, payload: dict) -> None:
    try:
        httpx.post(
            callback_url,
            json=payload,
            headers={"Authorization": f"Bearer {WORKER_SECRET}"},
            timeout=30,
        )
    except Exception as exc:
        print(f"[tailor_worker] callback failed: {exc}")


def _safe_filename(username: str, company: str) -> str:
    import re
    def clean(s: str) -> str:
        s = s.strip().lower()
        s = re.sub(r'[^\w\s-]', '', s)
        s = re.sub(r'[\s-]+', '_', s)
        return s[:40]
    return f"{clean(username)}_{clean(company)}"


def _compile_pdf_local(tex_path: Path) -> Path | None:
    """Try pdflatex locally; return PDF path on success, None otherwise."""
    if not shutil.which("pdflatex"):
        return None
    pdf_path = tex_path.with_suffix(".pdf")
    try:
        result = subprocess.run(
            ["pdflatex", "-interaction=nonstopmode",
             f"-output-directory={tex_path.parent}", str(tex_path)],
            capture_output=True, text=True, timeout=120,
        )
        if pdf_path.exists():
            print(f"[tailor_worker] PDF compiled: {pdf_path}")
            return pdf_path
        print(f"[tailor_worker] pdflatex failed: {result.stdout[-400:]}")
    except Exception as exc:
        print(f"[tailor_worker] pdflatex error: {exc}")
    return None


def _run_tailor_task(req: TailorJobRequest) -> None:
    """Run ResumeCrew on Hetzner and POST progress/results back to Render."""
    job_id = req.job_id

    def _cb(payload: dict) -> None:
        _post_tailor_callback(req.callback_url, {"job_id": job_id, **payload})

    def progress_callback(data: dict) -> None:
        _cb({"type": "progress", "event": data})

    try:
        # Ensure repo root is importable
        repo_root = str(_JOBEZEE_ROOT)
        if repo_root not in sys.path:
            sys.path.insert(0, repo_root)

        from PHASE2_JOB_TAILOR.crew import ResumeCrew

        job_dir = _JOBEZEE_ROOT / "job_outputs" / job_id
        job_dir.mkdir(parents=True, exist_ok=True)

        crew = ResumeCrew(openai_api_key=req.openai_api_key or None)
        result = crew.run_tailoring_process(
            req.job_description, req.resume_text, progress_callback, output_dir=job_dir
        )

        final_resume = result.get("final_resume", "")
        final_score  = result.get("final_score")
        latex_summary = result.get("latex_summary", {})
        company_raw  = result.get("keyword_analysis", {}).get("company_name", "") or req.company

        safe_name = _safe_filename(req.username, company_raw)

        tex_b64 = pdf_b64 = docx_b64 = ""

        # ── LaTeX → PDF ───────────────────────────────────────────────────────
        if latex_summary.get("status") == "success":
            src_tex = job_dir / "final_tailored_resume.tex"
            if src_tex.exists():
                dest_tex = job_dir / f"{safe_name}.tex"
                if src_tex != dest_tex:
                    shutil.copy2(src_tex, dest_tex)
                tex_b64 = base64.b64encode(dest_tex.read_bytes()).decode()
                pdf_file = _compile_pdf_local(dest_tex)
                if pdf_file:
                    pdf_b64 = base64.b64encode(pdf_file.read_bytes()).decode()

        # ── DOCX via python-docx ──────────────────────────────────────────────
        if final_resume:
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                docx_path = job_dir / f"{safe_name}.docx"
                doc = Document()
                for section in doc.sections:
                    section.top_margin = section.bottom_margin = Inches(0.6)
                    section.left_margin = section.right_margin = Inches(0.75)
                _HDRS = {"EXPERIENCE","EDUCATION","SKILLS","SUMMARY","OBJECTIVE",
                         "PROJECTS","CERTIFICATIONS","WORK EXPERIENCE","PROFESSIONAL EXPERIENCE"}
                first = True; second_done = False
                for line in final_resume.split("\n"):
                    s = line.strip()
                    if not s:
                        continue
                    if first:
                        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run(s); r.bold = True; r.font.size = Pt(18)
                        r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)
                        first = False; second_done = False; continue
                    if not second_done:
                        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run(s); r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0x44,0x44,0x44)
                        second_done = True; continue
                    is_hdr = (s.upper() in _HDRS or (s.isupper() and len(s)<50 and not s.startswith(('-','•','*'))))
                    if is_hdr:
                        p = doc.add_paragraph(style='Heading 2')
                        p.paragraph_format.space_before = Pt(8)
                        r = p.add_run(s.upper()); r.bold = True; r.font.size = Pt(11)
                        r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e); continue
                    if s.startswith(('-','•','*','·')):
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.left_indent = Inches(0.25)
                        r = p.add_run(s.lstrip('-•*· ').strip()); r.font.size = Pt(10); continue
                    p = doc.add_paragraph()
                    r = p.add_run(s); r.font.size = Pt(10)
                doc.save(str(docx_path))
                docx_b64 = base64.b64encode(docx_path.read_bytes()).decode()
            except Exception as exc:
                print(f"[tailor_worker] DOCX generation failed: {exc}")

        # ── Fallback PDF via reportlab ────────────────────────────────────────
        if not pdf_b64 and final_resume:
            try:
                from reportlab.lib.pagesizes import LETTER
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib import colors
                pdf_path = job_dir / f"{safe_name}.pdf"
                rdoc = SimpleDocTemplate(str(pdf_path), pagesize=LETTER,
                    leftMargin=0.75*inch, rightMargin=0.75*inch,
                    topMargin=0.75*inch, bottomMargin=0.75*inch)
                styles = getSampleStyleSheet()
                name_style = ParagraphStyle("Name", parent=styles["Normal"],
                    fontSize=16, fontName="Helvetica-Bold", spaceAfter=4,
                    textColor=colors.HexColor("#1a1a2e"))
                hdr_style = ParagraphStyle("Hdr", parent=styles["Normal"],
                    fontSize=11, fontName="Helvetica-Bold", spaceAfter=2,
                    textColor=colors.HexColor("#1a1a2e"))
                body_style = ParagraphStyle("Body", parent=styles["Normal"],
                    fontSize=10, fontName="Helvetica", leading=14, spaceAfter=1)
                _HDRS2 = {"EXPERIENCE","EDUCATION","SKILLS","SUMMARY","OBJECTIVE",
                          "PROJECTS","CERTIFICATIONS","AWARDS"}
                story = []; first2 = True
                for line in final_resume.split("\n"):
                    s = line.strip()
                    if not s: story.append(Spacer(1,4)); continue
                    if first2: story.append(Paragraph(s, name_style)); first2 = False; continue
                    if s.upper() in _HDRS2 or (len(s)<40 and s.isupper()):
                        story.append(Spacer(1,6)); story.append(Paragraph(s, hdr_style)); continue
                    safe_s = s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                    story.append(Paragraph(safe_s, body_style))
                rdoc.build(story)
                pdf_b64 = base64.b64encode(pdf_path.read_bytes()).decode()
            except Exception as exc:
                print(f"[tailor_worker] reportlab PDF failed: {exc}")

        _cb({
            "type":         "done",
            "status":       "complete",
            "score":        final_score,
            "final_resume": final_resume,
            "company_name": company_raw,
            "filename":     safe_name,
            "pdf_b64":      pdf_b64,
            "docx_b64":     docx_b64,
            "tex_b64":      tex_b64,
        })

    except Exception as exc:
        print(f"[tailor_worker] Job {job_id} failed: {exc}")
        _cb({"type": "done", "status": "error", "error": str(exc)})


@app.post("/run-tailor")
async def run_tailor(
    req: TailorJobRequest,
    background_tasks: BackgroundTasks,
    authorization: str = Header(...),
):
    """Accept a tailor job from Render, run ResumeCrew in background, POST results back."""
    _auth(authorization)
    background_tasks.add_task(_get_tailor_executor().submit, _run_tailor_task, req)
    return {"status": "queued", "job_id": req.job_id}


@app.get("/tailor-info")
def tailor_info(authorization: str = Header(...)):
    """Return tailor pool capacity."""
    _auth(authorization)
    ex = _get_tailor_executor()
    return {"max_workers": ex._max_workers}


# ── Compile PDF from LaTeX (existing callers in tailor_service.py) ────────────

class CompilePdfRequest(BaseModel):
    tex_b64:  str
    filename: str = "resume"


@app.post("/compile-pdf")
async def compile_pdf(req: CompilePdfRequest, authorization: str = Header(...)):
    """Compile a .tex file (base64) to PDF using pdflatex; return PDF as base64."""
    _auth(authorization)
    if not shutil.which("pdflatex"):
        raise HTTPException(503, "pdflatex not installed on this worker")
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        tex_path = tmp_path / f"{req.filename}.tex"
        tex_path.write_bytes(base64.b64decode(req.tex_b64))
        result = subprocess.run(
            ["pdflatex", "-interaction=nonstopmode",
             f"-output-directory={tmp}", str(tex_path)],
            capture_output=True, text=True, timeout=120,
        )
        pdf_path = tex_path.with_suffix(".pdf")
        if not pdf_path.exists():
            raise HTTPException(500, f"pdflatex failed: {result.stdout[-500:]}")
        pdf_b64 = base64.b64encode(pdf_path.read_bytes()).decode()
    return {"pdf_b64": pdf_b64}
