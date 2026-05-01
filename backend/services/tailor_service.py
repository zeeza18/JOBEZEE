"""
Tailor service — runs the PHASE2_JOB_TAILOR crew in a background thread
and tracks job state in memory.
"""
from __future__ import annotations

import os
import re
import shutil
import subprocess
import threading
import uuid
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any, Dict, Optional

import psutil

# PHASE2_JOB_TAILOR module root (two levels up from this file)
_SERVICE_DIR  = Path(__file__).resolve().parent          # backend/services/
_JOBEZEE_ROOT = _SERVICE_DIR.parent.parent               # JOBEZEE/
_TAILOR_ROOT  = _JOBEZEE_ROOT / "PHASE2_JOB_TAILOR"

# Ensure repo root is importable regardless of which service module loaded first
import sys as _sys
_JOBEZEE_ROOT_STR = str(_JOBEZEE_ROOT)
if _JOBEZEE_ROOT_STR not in _sys.path:
    _sys.path.insert(0, _JOBEZEE_ROOT_STR)
_OUTPUT_DIR   = _TAILOR_ROOT / "output"
_JOB_OUTPUTS  = _JOBEZEE_ROOT / "job_outputs"           # per-job output dirs


def _apply_contact_header(resume_text: str, contact_header: str) -> str:
    """
    Inject profile contact info (with real URLs) into the resume text.
    - If the name is NOT in the first line: prepend the full header.
    - If the name IS there (resume already has a header): replace the
      pipe-separated contact line so real LinkedIn/GitHub/Portfolio URLs
      overwrite the plain display words extracted from the original PDF.
    """
    if not contact_header:
        return resume_text
    header_lines = [l for l in contact_header.strip().split("\n") if l.strip()]
    name_hint = header_lines[0].strip() if header_lines else ""
    first_line = resume_text.strip().split("\n")[0].strip() if resume_text.strip() else ""

    if name_hint and name_hint.lower() not in first_line.lower():
        return contact_header + "\n\n" + resume_text

    # Name already present — replace the contact line (pipe-separated) with profile's
    profile_contact_line = next((l.strip() for l in header_lines if "|" in l), "")
    if not profile_contact_line:
        return resume_text

    r_lines = resume_text.split("\n")
    for i, line in enumerate(r_lines):
        if "|" in line and ("@" in line or any(
            x in line.lower() for x in ["linkedin", "github", "portfolio", "http", "+1"]
        )):
            r_lines[i] = profile_contact_line
            break
    return "\n".join(r_lines)

# In-memory job store (keyed by tailor_job_id)
_jobs: Dict[str, Dict[str, Any]] = {}
_lock = threading.Lock()

# Thread pool — size based on available RAM
def _max_workers() -> int:
    return 2  # 2 concurrent tailor slots; extras queue automatically

_executor: Optional[ThreadPoolExecutor] = None
_executor_lock = threading.Lock()

def _get_executor() -> ThreadPoolExecutor:
    global _executor
    with _executor_lock:
        if _executor is None or _executor._shutdown:
            n = _max_workers()
            _executor = ThreadPoolExecutor(max_workers=n, thread_name_prefix="tailor")
            print(f"[tailor_service] Worker pool: {n} concurrent jobs")
    return _executor


def get_max_workers() -> int:
    """Return the current pool capacity (RAM-based)."""
    return _get_executor()._max_workers


def get_active_count() -> int:
    """Return the number of jobs currently executing."""
    with _lock:
        return sum(1 for j in _jobs.values() if j.get("status") == "running")


def create_job(user_id: str = "") -> str:
    job_id = str(uuid.uuid4())
    with _lock:
        _jobs[job_id] = {
            "user_id":      user_id,
            "status":       "queued",
            "progress":     [],
            "score":        None,
            "company_name": None,
            "tex_path":     None,
            "pdf_path":     None,
            "docx_path":    None,
            "final_resume": None,
            "filename":     None,
            "error":        None,
        }
    return job_id


def get_job(job_id: str) -> Optional[Dict[str, Any]]:
    return _jobs.get(job_id)


def get_user_jobs(user_id: str) -> list:
    """Return all in-memory jobs for a user (newest first by insertion order)."""
    with _lock:
        return [
            {"job_id": jid, **jdata}
            for jid, jdata in reversed(list(_jobs.items()))
            if jdata.get("user_id") == user_id
        ]


# ─── Original: run from plain text (TailorPage) ───────────────────────────────

def _hetzner_url() -> str:
    return os.getenv("BOT_WORKER_URL", "").rstrip("/")

def _worker_secret() -> str:
    return os.getenv("WORKER_SECRET", "")

def _api_base_url() -> str:
    return os.getenv("API_BASE_URL", "https://www.jobezee.org")

def _delegate_to_hetzner(
    job_id: str,
    job_description: str,
    resume_text: str,
    openai_api_key: str,
    username: str,
    company: str,
    anthropic_api_key: str = "",
    anthropic_fallback_key: str = "",
) -> None:
    """POST tailor job to Hetzner worker asynchronously (fire-and-forget thread)."""
    import httpx as _httpx, threading as _thr

    payload = {
        "job_id":                  job_id,
        "job_description":         job_description,
        "resume_text":             resume_text,
        "openai_api_key":          openai_api_key,
        "anthropic_api_key":       anthropic_api_key,
        "anthropic_fallback_key":  anthropic_fallback_key,
        "username":                username,
        "company":                 company,
        "callback_url":            f"{_api_base_url()}/api/tailor/internal/callback",
    }

    def _send():
        try:
            r = _httpx.post(
                f"{_hetzner_url()}/run-tailor",
                json=payload,
                headers={"Authorization": f"Bearer {_worker_secret()}"},
                timeout=30,
            )
            print(f"[tailor_service] Hetzner queued job {job_id}: {r.status_code}")
        except Exception as exc:
            print(f"[tailor_service] Hetzner delegation failed for {job_id}: {exc}")
            with _lock:
                _jobs[job_id]["status"] = "error"
                _jobs[job_id]["error"]  = f"Worker unreachable: {exc}"

    _thr.Thread(target=_send, daemon=True).start()


def start_tailor_job(job_id: str, job_description: str, resume: str, openai_api_key: str = "", username: str = "", anthropic_api_key: str = "", anthropic_fallback_key: str = "") -> None:
    """Submit job — delegates to Hetzner if BOT_WORKER_URL is set, else local thread pool."""
    if _hetzner_url():
        with _lock:
            _jobs[job_id]["status"] = "running"
        _delegate_to_hetzner(job_id, _clean_job_description(job_description), resume, openai_api_key, username, "", anthropic_api_key, anthropic_fallback_key)
    else:
        _get_executor().submit(_run_tailor_job, job_id, job_description, resume, openai_api_key, username, anthropic_api_key, anthropic_fallback_key)


def _run_tailor_job(job_id: str, job_description: str, resume: str, openai_api_key: str = "", username: str = "", anthropic_api_key: str = "", anthropic_fallback_key: str = "") -> None:
    """Runs inside background thread."""
    import sys as _sys
    with _lock:
        _jobs[job_id]["status"] = "running"

    def progress_callback(data: Dict[str, Any]) -> None:
        with _lock:
            _jobs[job_id]["progress"].append(data)
            if data.get("event") == "round_complete":
                score = data.get("evaluation", {}).get("score")
                if score is not None:
                    _jobs[job_id]["score"] = score

    def _cp1252_safe(s: str) -> str:
        if _sys.platform != "win32":
            return s  # Linux/Render: UTF-8 default, no stripping needed
        out = []
        for c in (s or ""):
            try:
                c.encode('cp1252')
                out.append(c)
            except (UnicodeEncodeError, LookupError):
                out.append(' ')
        return ''.join(out)

    try:
        _clean_jd = _cp1252_safe(_clean_job_description(job_description))
        _resume   = _cp1252_safe(resume)

        from PHASE2_JOB_TAILOR.crew import ResumeCrew

        job_dir = _JOB_OUTPUTS / job_id
        job_dir.mkdir(parents=True, exist_ok=True)

        crew = ResumeCrew(openai_api_key=openai_api_key or None, anthropic_api_key=anthropic_api_key or None, anthropic_fallback_key=anthropic_fallback_key or None)
        result = crew.run_tailoring_process(_clean_jd, _resume, progress_callback, output_dir=job_dir)

        final_score = result.get("final_score")
        final_resume = result.get("final_resume", "")
        latex_summary = result.get("latex_summary", {})

        tex_path: Optional[Path] = None
        pdf_path: Optional[Path] = None

        # Build filename: username_company (from Tool 1 keyword_analysis)
        company_raw = result.get("keyword_analysis", {}).get("company_name", "") or "company"
        with _lock:
            _jobs[job_id]["company_name"] = company_raw
        progress_callback({"event": "company_detected", "company_name": company_raw})
        safe_name = _safe_filename(username or "resume", company_raw)

        if latex_summary.get("status") == "success":
            candidate = job_dir / "final_tailored_resume.tex"
            if candidate.exists():
                dest_tex = job_dir / f"{safe_name}.tex"
                if candidate != dest_tex:
                    shutil.copy2(candidate, dest_tex)
                tex_path = dest_tex
                compiled = _compile_pdf(dest_tex)
                if compiled:
                    pdf_path = compiled

        if not pdf_path and final_resume:
            pdf_path = _generate_pdf_from_text(final_resume, job_dir / f"{safe_name}.pdf")

        docx_path: Optional[Path] = None
        if final_resume:
            docx_path = _generate_docx_from_text(final_resume, job_dir / f"{safe_name}.docx")

        with _lock:
            _jobs[job_id].update({
                "status": "complete",
                "score": final_score,
                "final_resume": final_resume,
                "tex_path": str(tex_path) if tex_path else None,
                "pdf_path": str(pdf_path) if pdf_path else None,
                "docx_path": str(docx_path) if docx_path else None,
                "filename": safe_name,
            })

    except Exception as exc:
        with _lock:
            _jobs[job_id].update({
                "status": "error",
                "error": str(exc),
            })
        print(f"[tailor_service] Job {job_id} failed: {exc}")


# ─── New: run from uploaded resume file (job-card Tailor button) ─────────────

def start_tailor_job_for_job(
    tailor_job_id: str,
    job_description: str,
    resume_text: str,
    username: str,
    company: str,
    openai_api_key: str = "",
    contact_header: str = "",
    anthropic_api_key: str = "",
    anthropic_fallback_key: str = "",
) -> None:
    """Submit job — delegates to Hetzner if BOT_WORKER_URL is set, else local thread pool."""
    if _hetzner_url():
        import threading as _thr

        def _prepare_and_delegate():
            try:
                final_resume = _apply_contact_header(resume_text, contact_header)

                clean_jd = _clean_job_description(job_description)
                raw_words   = len(job_description.split())
                clean_words = len(clean_jd.split())
                if clean_words < 40 and raw_words >= clean_words:
                    import re as _re
                    clean_jd = _re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', job_description)
                    clean_jd = _re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', clean_jd)
                if len(clean_jd.split()) < 10:
                    raise RuntimeError("Job description too short to tailor against.")

                with _lock:
                    _jobs[tailor_job_id]["status"] = "running"

                _delegate_to_hetzner(tailor_job_id, clean_jd, final_resume, openai_api_key, username, company, anthropic_api_key, anthropic_fallback_key)
            except Exception as exc:
                with _lock:
                    _jobs[tailor_job_id]["status"] = "error"
                    _jobs[tailor_job_id]["error"]  = str(exc)

        _thr.Thread(target=_prepare_and_delegate, daemon=True).start()
    else:
        _get_executor().submit(_run_tailor_for_job, tailor_job_id, job_description, resume_text, username, company, openai_api_key, contact_header, anthropic_api_key, anthropic_fallback_key)


def _run_tailor_for_job(
    tailor_job_id: str,
    job_description: str,
    resume_text: str,
    username: str,
    company: str,
    openai_api_key: str = "",
    contact_header: str = "",
    anthropic_api_key: str = "",
    anthropic_fallback_key: str = "",
) -> None:
    with _lock:
        _jobs[tailor_job_id]["status"] = "running"
        _jobs[tailor_job_id]["company_name"] = company

    def progress_callback(data: Dict[str, Any]) -> None:
        with _lock:
            _jobs[tailor_job_id]["progress"].append(data)
            if data.get("event") == "round_complete":
                score = data.get("evaluation", {}).get("score")
                if score is not None:
                    _jobs[tailor_job_id]["score"] = score

    # Emit company_detected immediately since company is already known
    progress_callback({"event": "company_detected", "company_name": company})

    try:
        resume_text = _apply_contact_header(resume_text, contact_header)
        clean_jd = _clean_job_description(job_description)

        # If cleaning removed too much (boilerplate was all we had), fall back to
        # raw JD with only markdown stripped — better than erroring or feeding nothing.
        raw_words = len(job_description.split())
        clean_words = len(clean_jd.split())
        if clean_words < 40 and raw_words >= clean_words:
            print(f"[tailor_service] Cleaned JD too short ({clean_words}w), using raw JD ({raw_words}w)")
            clean_jd = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', job_description)
            clean_jd = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', clean_jd)

        if len(clean_jd.split()) < 10:
            raise RuntimeError(
                "Job description is too short to tailor against — "
                "the scraper only captured a preview of this posting. "
                "Pull fresh jobs to get the full description."
            )

        # Sanitize chars Windows cp1252 can't encode (arrows, emojis, curly chars)
        # Only needed on Windows — Linux/Render uses UTF-8 natively
        import sys as _sys
        if _sys.platform == "win32":
            def _cp1252_safe(s: str) -> str:
                out = []
                for c in (s or ""):
                    try:
                        c.encode('cp1252')
                        out.append(c)
                    except (UnicodeEncodeError, LookupError):
                        out.append(' ')
                return ''.join(out)
            clean_jd    = _cp1252_safe(clean_jd)
            resume_text = _cp1252_safe(resume_text)

        from PHASE2_JOB_TAILOR.crew import ResumeCrew

        safe_name = _safe_filename(username, company)
        job_dir = _JOB_OUTPUTS / tailor_job_id
        job_dir.mkdir(parents=True, exist_ok=True)

        crew = ResumeCrew(openai_api_key=openai_api_key or None, anthropic_api_key=anthropic_api_key or None, anthropic_fallback_key=anthropic_fallback_key or None)
        result = crew.run_tailoring_process(clean_jd, resume_text, progress_callback, output_dir=job_dir)

        final_score = result.get("final_score")
        final_resume = result.get("final_resume", "")
        latex_summary = result.get("latex_summary", {})

        tex_path: Optional[Path] = None
        pdf_path: Optional[Path] = None
        docx_path: Optional[Path] = None

        if latex_summary.get("status") == "success":
            src_tex = job_dir / "final_tailored_resume.tex"
            if src_tex.exists():
                dest_tex = job_dir / f"{safe_name}.tex"
                if src_tex != dest_tex:
                    shutil.copy2(src_tex, dest_tex)
                tex_path = dest_tex
                compiled = _compile_pdf(dest_tex)
                if compiled:
                    pdf_path = compiled

        if not pdf_path and final_resume:
            pdf_path = _generate_pdf_from_text(final_resume, job_dir / f"{safe_name}.pdf")

        if final_resume:
            docx_path = _generate_docx_from_text(final_resume, job_dir / f"{safe_name}.docx")

        with _lock:
            _jobs[tailor_job_id].update({
                "status": "complete",
                "score": final_score,
                "final_resume": final_resume,
                "tex_path": str(tex_path) if tex_path else None,
                "pdf_path": str(pdf_path) if pdf_path else None,
                "docx_path": str(docx_path) if docx_path else None,
                "filename": safe_name,
            })

    except Exception as exc:
        with _lock:
            _jobs[tailor_job_id].update({
                "status": "error",
                "error": str(exc),
            })
        print(f"[tailor_service] Job {tailor_job_id} failed: {exc}")


def _extract_resume_text(file_path: Path) -> str:
    """Extract plain text from PDF or DOCX resume file."""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
            return "\n\n".join(p for p in pages if p.strip())
        except ImportError:
            print("[tailor_service] pdfplumber not installed — trying PyPDF2")
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return "\n\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
        except ImportError:
            raise RuntimeError(
                "PDF extraction requires pdfplumber. Run: pip install pdfplumber"
            )

    elif suffix in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except ImportError:
            raise RuntimeError(
                "DOCX extraction requires python-docx. Run: pip install python-docx"
            )

    elif suffix == ".txt":
        return file_path.read_text(encoding="utf-8", errors="ignore")

    else:
        # Try reading as plain text as last resort
        try:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            raise RuntimeError(f"Unsupported resume format: {suffix}")


def _clean_job_description(jd: str) -> str:
    """
    Extract only the requirements-relevant sections from a scraped job description.

    Strategy (keeper-only mode):
    - Strip markdown formatting
    - DEFAULT = skip.  Only include lines that fall under a recognised KEEPER
      section header (responsibilities, skills, requirements, qualifications, etc.)
    - A KEEPER header is included in output; everything before the first KEEPER
      header (intro / "About the job" blurb) is silently dropped
    - A BOILERPLATE header (company overview, benefits, EEO, …) exits keeper mode
    - Unknown standalone headers also exit keeper mode so unrecognised sections
      don't accidentally pull in salary/perks/location prose
    - Trailing EEO paragraphs stripped regardless of section structure
    - If the result is too short the caller falls back to the raw JD

    This keeps only what the tailor AI actually needs, cutting token waste.
    """
    if not jd:
        return jd

    # ── 1. Strip markdown formatting ─────────────────────────────────────────
    text = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', jd)
    text = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', text)
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)

    # ── 2. KEEPER section headers — enter include mode ───────────────────────
    # Broad set of patterns covering many JD styles
    _KEEPER = re.compile(
        r'^(?:'
        # responsibilities
        r'responsibilities?'
        r'|key\s+responsibilities?'
        r'|core\s+responsibilities?'
        r'|primary\s+responsibilities?'
        r'|main\s+responsibilities?'
        r'|what\s+you.ll\s+(do|build|work\s+on|own|own\s+and\s+build|be\s+doing)'
        r'|what\s+you\s+will\s+do'
        r'|your\s+(role|impact|responsibilities?|day[\s-]to[\s-]day)'
        r'|day[\s-]to[\s-]day'
        r'|you\s+will'
        r'|in\s+this\s+role'
        # requirements / qualifications
        r'|requirements?'
        r'|qualifications?'
        r'|required\s+qualifications?'
        r'|preferred\s+qualifications?'
        r'|minimum\s+qualifications?'
        r'|basic\s+qualifications?'
        r'|desired\s+qualifications?'
        r'|technical\s+requirements?'
        r'|position\s+requirements?'
        r'|job\s+requirements?'
        r'|what\s+(you|we).re?\s+looking\s+for'
        r'|what\s+you.ll\s+bring'
        r'|what\s+you\s+bring'
        r'|what\s+we\s+need'
        r'|what\s+you\s+need'
        r'|must[\s-]have'
        r'|nice[\s-]to[\s-]have'
        r'|we\s+are\s+looking'
        # skills
        r'|skills?\s*(required|needed|we\s+need|&\s+experience)?'
        r'|required\s+skills?'
        r'|technical\s+skills?'
        r'|core\s+skills?'
        r'|key\s+skills?'
        r'|essential\s+skills?'
        r'|desired\s+skills?'
        r'|preferred\s+skills?'
        # experience / education
        r'|experience'
        r'|required\s+experience'
        r'|preferred\s+experience'
        r'|relevant\s+experience'
        r'|education'
        r'|certifications?'
        # tech stack
        r'|tech\s+(stack|skills?|requirements?)'
        r'|technologies'
        r'|tools?\s+(&\s+technologies?)?'
        r'|stack'
        # LinkedIn structured block
        r'|requirements\s+added\s+by.*'
        r')\s*[:\-]?\s*$',
        re.IGNORECASE,
    )

    # ── 3. BOILERPLATE section headers — exit include mode ───────────────────
    _BOILERPLATE = re.compile(
        r'^(?:'
        r'about\s+(us|the\s+company|the\s+team|our\s+company|the\s+job|the\s+role|the\s+position|[a-z]{2,30})\b'
        r'|who\s+we\s+are'
        r'|our\s+(story|mission|values?|culture|vision|team)'
        r'|company\s+overview'
        r'|overview'
        r'|introduction'
        r'|why\s+(join\s+us|work\s+(here|with\s+us|for\s+us)|us|choose\s+us)'
        r'|what\s+we\s+offer'
        r'|what\s+makes\s+this\s+role\s+(different|unique|special|stand\s+out)'
        r'|what\s+sets\s+(us|this\s+role)\s+apart'
        r'|life\s+at\s+\w+'
        r'|compensation(\s+and\s+benefits?)?'
        r'|benefits?\s+(&|and)\s+perks?'
        r'|perks?\s+(&|and)\s+benefits?'
        r'|benefits?\s*$'
        r'|perks?\s*$'
        r'|featured\s+benefits?'
        r'|what\s+you.ll\s+(get|receive|enjoy)'
        r'|additional\s+details?'
        r'|additional\s+information'
        r'|interview\s+(process|stages?|steps?)'
        r'|our\s+interview\s+process'
        r'|hiring\s+process'
        r'|how\s+we\s+(hire|interview|work)'
        r'|equal\s+(opportunity|employment)'
        r'|eeo\b'
        r'|diversity\s+(&|and)\s+inclusion'
        r'|we\s+are\s+an\s+equal'
        r'|salary\s+range'
        r'|pay\s+range'
        r'|total\s+(rewards?|compensation)'
        r'|application\s+(terms|process|note|deadline)'
        r'|please\s+(note|read)\b'
        r'|note\s+to\s+(applicants?|candidates?)'
        r'|working\s+(conditions?|environment)'
        r'|work\s+environment'
        r'|location'
        r'|office\s+(location|hours?)'
        r')\s*[:\-]?\s*$',
        re.IGNORECASE,
    )

    lines = text.split('\n')
    output_lines: list[str] = []
    in_keeper = False   # start in SKIP mode — only include known keeper sections

    for line in lines:
        stripped = line.strip()

        # A "section header" = standalone short line, not a bullet or long sentence
        is_header = (
            bool(stripped)
            and len(stripped) < 80
            and not stripped.startswith(('-', '*', '•', '·', '–', '◦'))
            and not stripped[0].isdigit()   # numbered bullets like "1. Do X"
        )

        if is_header and _KEEPER.match(stripped):
            in_keeper = True
            output_lines.append(line)
            continue

        if is_header and _BOILERPLATE.match(stripped):
            in_keeper = False
            continue

        if is_header:
            # Unknown standalone header — keep current keeper state.
            # If we're already in a keeper section, include the line (it may be a
            # relevant sub-header like "Technologies Used:" not yet in our list).
            # If we're not in a keeper section, drop it (pre-intro prose, etc.).
            if in_keeper:
                output_lines.append(line)
            continue

        if in_keeper:
            output_lines.append(line)

    cleaned = '\n'.join(output_lines)

    # ── 4. Collapse excessive blank lines ────────────────────────────────────
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)

    # ── 5. Strip trailing EEO / application-note paragraphs (up to last 15 lines)
    _EEO_PHRASES = [
        'equal opportunity', 'eeo', 'affirmative action', 'discrimination',
        'applicants will receive consideration', 'regardless of race',
        'regardless of gender', 'protected veteran', 'disability status',
        'visa sponsorship', 'authorized to work', 'work authorization',
        'unsubscribe', 'resume database', 'resume submission',
        'background check', 'drug test',
    ]
    final_lines = cleaned.rstrip().split('\n')
    removed = 0
    while final_lines and removed < 15:
        last = final_lines[-1].lower()
        if any(phrase in last for phrase in _EEO_PHRASES):
            final_lines.pop()
            removed += 1
        else:
            break
    cleaned = '\n'.join(final_lines).strip()

    print(f"[tailor_service] JD cleaned: {len(jd)} -> {len(cleaned)} chars")
    return cleaned


def _safe_filename(username: str, company: str) -> str:
    """Produce a filesystem-safe filename like 'john_doe_google'."""
    def clean(s: str) -> str:
        s = s.strip().lower()
        s = re.sub(r'[^\w\s-]', '', s)
        s = re.sub(r'[\s-]+', '_', s)
        return s[:40]
    return f"{clean(username)}_{clean(company)}"


def _generate_docx_from_text(resume_text: str, out_path: Path) -> Optional[Path]:
    """Generate a properly formatted Word .docx from plain-text resume using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Narrow margins
        for section in doc.sections:
            section.top_margin    = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin   = Inches(0.75)
            section.right_margin  = Inches(0.75)

        _SECTION_HEADERS = {
            "EXPERIENCE", "EDUCATION", "SKILLS", "SUMMARY", "OBJECTIVE",
            "PROJECTS", "CERTIFICATIONS", "AWARDS", "PUBLICATIONS",
            "VOLUNTEER", "LANGUAGES", "INTERESTS", "REFERENCES",
            "TECHNICAL SKILLS", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE",
        }

        lines = resume_text.split("\n")
        first_non_empty = True
        second_line_done = False

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # First line = Name (large, bold, centered)
            if first_non_empty:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
                first_non_empty = False
                second_line_done = False
                continue

            # Second line = contact info (centered, smaller)
            if not second_line_done:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                second_line_done = True
                continue

            # Section headers (ALL CAPS short lines)
            is_header = (stripped.upper() in _SECTION_HEADERS or
                         (stripped.isupper() and len(stripped) < 50 and not stripped.startswith(('-', '•', '*'))))
            if is_header:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(stripped.upper())
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
                # Add underline via border — approximate with a horizontal rule style
                p.style = doc.styles['Heading 2']
                run2 = p.runs[0] if p.runs else p.add_run(stripped.upper())
                run2.bold = True
                run2.font.size = Pt(11)
                run2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
                continue

            # Bullet points
            if stripped.startswith(('-', '•', '*', '·')):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(stripped.lstrip('-•*· ').strip())
                run.font.size = Pt(10)
                continue

            # Everything else = normal paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped)
            run.font.size = Pt(10)

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        print(f"[tailor_service] DOCX generated: {out_path}")
        return out_path
    except ImportError:
        print("[tailor_service] python-docx not installed — DOCX generation skipped")
    except Exception as exc:
        print(f"[tailor_service] DOCX generation failed: {exc}")
    return None


def _generate_pdf_from_text(resume_text: str, out_path: Path) -> Optional[Path]:
    """Generate a clean PDF from plain-text resume using reportlab (pdflatex fallback)."""
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib import colors

        doc = SimpleDocTemplate(
            str(out_path),
            pagesize=LETTER,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )
        styles = getSampleStyleSheet()
        heading_style = ParagraphStyle(
            "Heading",
            parent=styles["Normal"],
            fontSize=11,
            fontName="Helvetica-Bold",
            spaceAfter=2,
            textColor=colors.HexColor("#1a1a2e"),
        )
        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontSize=10,
            fontName="Helvetica",
            leading=14,
            spaceAfter=1,
        )
        name_style = ParagraphStyle(
            "Name",
            parent=styles["Normal"],
            fontSize=16,
            fontName="Helvetica-Bold",
            spaceAfter=4,
            textColor=colors.HexColor("#1a1a2e"),
        )

        story = []
        _SECTION_HEADERS = {
            "EXPERIENCE", "EDUCATION", "SKILLS", "SUMMARY", "OBJECTIVE",
            "PROJECTS", "CERTIFICATIONS", "AWARDS", "PUBLICATIONS",
            "VOLUNTEER", "LANGUAGES", "INTERESTS", "REFERENCES",
        }

        lines = resume_text.split("\n")
        first_non_empty = True
        for line in lines:
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 4))
                continue
            # First non-empty line = name
            if first_non_empty:
                story.append(Paragraph(stripped, name_style))
                first_non_empty = False
            elif stripped.upper() in _SECTION_HEADERS or (len(stripped) < 40 and stripped.isupper()):
                story.append(Spacer(1, 6))
                story.append(Paragraph(stripped, heading_style))
            else:
                # Escape HTML special chars for reportlab
                safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, body_style))

        doc.build(story)
        print(f"[tailor_service] PDF generated via reportlab: {out_path}")
        return out_path
    except ImportError:
        print("[tailor_service] reportlab not installed — PDF generation skipped")
    except Exception as exc:
        print(f"[tailor_service] reportlab PDF generation failed: {exc}")
    return None


def _compile_pdf(tex_path: Path) -> Optional[Path]:
    """Run pdflatex locally, or delegate to Hetzner worker if not installed here."""
    import base64 as _b64, os as _os

    pdf_path = tex_path.with_suffix(".pdf")

    # ── Try local pdflatex first ──────────────────────────────────────────────
    if shutil.which("pdflatex"):
        try:
            result = subprocess.run(
                ["pdflatex", "-interaction=nonstopmode",
                 f"-output-directory={tex_path.parent}", str(tex_path)],
                capture_output=True, text=True, timeout=120,
            )
            if pdf_path.exists():
                print(f"[tailor_service] PDF compiled locally: {pdf_path}")
                return pdf_path
            print(f"[tailor_service] pdflatex failed: {result.stdout[-400:]}")
        except Exception as exc:
            print(f"[tailor_service] local pdflatex error: {exc}")

    # ── Delegate to Hetzner worker ────────────────────────────────────────────
    worker_url    = _os.getenv("BOT_WORKER_URL", "").rstrip("/")
    worker_secret = _os.getenv("WORKER_SECRET", "")
    if worker_url and worker_secret and tex_path.exists():
        try:
            import httpx as _httpx
            tex_b64  = _b64.b64encode(tex_path.read_bytes()).decode()
            filename = tex_path.stem
            r = _httpx.post(
                f"{worker_url}/compile-pdf",
                json={"tex_b64": tex_b64, "filename": filename},
                headers={"Authorization": f"Bearer {worker_secret}"},
                timeout=150,
            )
            if r.status_code == 200:
                pdf_bytes = _b64.b64decode(r.json()["pdf_b64"])
                pdf_path.write_bytes(pdf_bytes)
                print(f"[tailor_service] PDF compiled via Hetzner: {pdf_path} ({len(pdf_bytes)} bytes)")
                return pdf_path
            else:
                print(f"[tailor_service] Hetzner compile-pdf error: {r.status_code} {r.text[:200]}")
        except Exception as exc:
            print(f"[tailor_service] Hetzner compile-pdf failed: {exc}")

    print("[tailor_service] pdflatex not available locally and Hetzner delegation failed/skipped")
    return None
