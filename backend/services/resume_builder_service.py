"""
Resume Maker service — structured resume documents, AI-assisted import/scoring,
and PDF/DOCX export.

Distinct from tailor_service.py: Tailor rewrites an existing flat-text resume
against a job description. Maker builds a structured, template-driven resume
from scratch (or from an imported one) and exports it.
"""
from __future__ import annotations

import base64
import json
import re
import tempfile
from pathlib import Path
from typing import Any

import httpx
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from ..config import get_settings
from ..models import ResumeDocument, UserProfile
from ..schemas import (
    ResumeDocumentContent,
    ResumeDocumentSettings,
    ResumeDocumentUpdate,
    ResumeScoreItemFeedback,
    ResumeScoreResponse,
)
from .tailor_service import _extract_resume_text

OPENAI_CHAT_URL = "https://api.openai.com/v1/chat/completions"
OPENAI_MODEL    = "gpt-4o-mini"


# ---------------------------------------------------------------------------
# Document CRUD — a user can have any number of resumes
# ---------------------------------------------------------------------------

async def list_documents(db: AsyncSession, user_id: str) -> list[ResumeDocument]:
    result = await db.execute(
        select(ResumeDocument)
        .where(ResumeDocument.user_id == user_id)
        .order_by(ResumeDocument.updated_at.desc())
    )
    return list(result.scalars().all())


async def get_document(db: AsyncSession, user_id: str, doc_id: str) -> ResumeDocument:
    """Fetch a document, scoped to its owner. Raises LookupError if missing/not owned."""
    result = await db.execute(
        select(ResumeDocument).where(ResumeDocument.id == doc_id, ResumeDocument.user_id == user_id)
    )
    doc = result.scalar_one_or_none()
    if doc is None:
        raise LookupError("Resume document not found")
    return doc


async def create_document(
    db: AsyncSession, user_id: str, title: str = "My Resume", seed_from_profile: bool = False,
) -> ResumeDocument:
    """Create a new document. seed_from_profile pre-fills it from the user's uploaded
    resume (used only for a user's very first document) — falls back to blank on failure."""
    content = ResumeDocumentContent()
    if seed_from_profile:
        try:
            content = await import_from_profile(db, user_id)
        except Exception:
            pass
    doc = ResumeDocument(
        user_id=user_id,
        title=title,
        content=content.model_dump(mode="json"),
        settings=ResumeDocumentSettings().model_dump(mode="json"),
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return doc


async def duplicate_document(db: AsyncSession, user_id: str, doc_id: str) -> ResumeDocument:
    src = await get_document(db, user_id, doc_id)
    doc = ResumeDocument(
        user_id=user_id,
        title=f"{src.title} (copy)",
        content=src.content,
        settings=src.settings,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return doc


async def update_document(db: AsyncSession, user_id: str, doc_id: str, update: ResumeDocumentUpdate) -> ResumeDocument:
    doc = await get_document(db, user_id, doc_id)
    if update.title is not None:
        doc.title = update.title
    if update.content is not None:
        doc.content = update.content.model_dump(mode="json")
    if update.settings is not None:
        doc.settings = update.settings.model_dump(mode="json")
    await db.commit()
    await db.refresh(doc)
    return doc


async def delete_document(db: AsyncSession, user_id: str, doc_id: str) -> None:
    doc = await get_document(db, user_id, doc_id)
    await db.delete(doc)
    await db.commit()


# ---------------------------------------------------------------------------
# OpenAI helper (house pattern — see routers/avatar_interview.py)
# ---------------------------------------------------------------------------

async def _openai_chat_json(system: str, user: str, max_tokens: int = 2000) -> dict[str, Any]:
    cfg = get_settings()
    api_key = (cfg.OPENAI_API_KEY or "").strip()
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY not configured on server")

    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.post(
            OPENAI_CHAT_URL,
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": OPENAI_MODEL,
                "max_tokens": max_tokens,
                "temperature": 0.2,
                "response_format": {"type": "json_object"},
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
            },
        )
    resp.raise_for_status()
    raw = resp.json()["choices"][0]["message"]["content"]
    cleaned = re.sub(r"```json\s*|```\s*", "", raw).strip()
    return json.loads(cleaned)


# ---------------------------------------------------------------------------
# Import from profile (auto-structuring — the "do better than Resume-Matcher" bit)
# ---------------------------------------------------------------------------

async def _get_profile_resume_text(db: AsyncSession, user_id: str) -> str:
    """Same file-resolution logic as GET /api/tailor/profile-resume-text
    (resume_url on disk, falling back to DB-stored base64 bytes)."""
    import uuid as _uuid

    cfg = get_settings()
    pid = _uuid.UUID(user_id)
    result = await db.execute(select(UserProfile).where(UserProfile.id == pid))
    profile = result.scalar_one_or_none()
    resume_url = (profile.resume_url if profile else "") or ""
    if not resume_url:
        raise RuntimeError("No resume on file. Upload one via Profile → Resume first.")

    jobezee_root = Path(__file__).resolve().parent.parent.parent
    upload_root = jobezee_root / cfg.UPLOAD_DIR
    rel = resume_url.lstrip("/")
    if rel.startswith("uploads/"):
        rel = rel[len("uploads/"):]
    file_path = upload_root / rel

    if not file_path.exists():
        b64 = getattr(profile, "resume_bytes", "") or ""
        if not b64:
            raise RuntimeError(f"Resume file not found on server ({file_path.name}). Please re-upload.")
        tmp_dir = Path(tempfile.mkdtemp())
        tmp_path = tmp_dir / file_path.name
        tmp_path.write_bytes(base64.b64decode(b64))
        file_path = tmp_path

    suffix = file_path.suffix.lower()
    text = file_path.read_text(errors="replace") if suffix == ".txt" else _extract_resume_text(file_path)
    if not text.strip():
        raise RuntimeError("Could not extract text from the stored resume.")
    return text


_IMPORT_SYSTEM_PROMPT = """You convert a plain-text resume into structured JSON for a resume builder.
Respond with JSON ONLY (no markdown fences), matching exactly this shape:
{
  "contact": {"full_name": "", "headline": "", "email": "", "phone": "", "location": "", "linkedin": "", "github": "", "portfolio": "", "website": ""},
  "summary": "",
  "experience": [{"id": "", "company": "", "title": "", "location": "", "start_date": "", "end_date": "", "current": false, "bullets": [""]}],
  "education": [{"id": "", "school": "", "degree": "", "field": "", "location": "", "start_date": "", "end_date": "", "gpa": ""}],
  "skills": [{"id": "", "label": "", "items": [""]}],
  "projects": [{"id": "", "name": "", "description": "", "bullets": [""], "link": "", "tech": [""]}],
  "certifications": [{"id": "", "name": "", "issuer": "", "date": "", "link": ""}]
}
Preserve every fact verbatim (company names, dates, metrics, titles) — never invent or embellish content.
Give each list item a short unique "id" slug (e.g. "exp-1", "edu-1"). Split the skills into 1-3 sensible
categories (e.g. "Languages", "Frameworks", "Tools"). Omit section_order (the client fills it in)."""


async def import_from_profile(db: AsyncSession, user_id: str) -> ResumeDocumentContent:
    resume_text = await _get_profile_resume_text(db, user_id)
    data = await _openai_chat_json(_IMPORT_SYSTEM_PROMPT, resume_text[:12000], max_tokens=3000)
    return ResumeDocumentContent.model_validate(data)


# ---------------------------------------------------------------------------
# ATS-style scoring (OpenAI — live panel in the Maker preview)
# ---------------------------------------------------------------------------

def _flatten_content(content: ResumeDocumentContent) -> str:
    lines: list[str] = []
    c = content.contact
    lines.append(f"{c.full_name} — {c.headline}")
    if content.summary:
        lines.append(content.summary)
    for exp in content.experience:
        lines.append(f"{exp.title} at {exp.company} ({exp.start_date}–{exp.end_date or 'Present'})")
        lines.extend(f"- {b}" for b in exp.bullets)
    for edu in content.education:
        lines.append(f"{edu.degree} {edu.field}, {edu.school}")
    for sk in content.skills:
        lines.append(f"{sk.label}: {', '.join(sk.items)}")
    for proj in content.projects:
        lines.append(f"Project: {proj.name} — {proj.description}")
        lines.extend(f"- {b}" for b in proj.bullets)
    for cert in content.certifications:
        lines.append(f"Certification: {cert.name} ({cert.issuer})")
    return "\n".join(lines)


_SCORE_SYSTEM_PROMPT = """You are a strict ATS (applicant tracking system) resume reviewer.
Respond with JSON ONLY (no markdown fences), matching exactly this shape:
{
  "score": 0,
  "matched": [""],
  "missing": [""],
  "suggestions": [""],
  "item_feedback": [{"section": "experience", "snippet": "first 6-10 words of the exact bullet/line", "issue": "what's wrong and how to fix it"}]
}
"score" is 0-100 overall resume quality/ATS-parseability (and job-fit if a job description is given).
"matched" and "missing" are notable keywords/skills present vs absent (job description keywords if given,
otherwise general strong-resume signals). "suggestions" is up to 5 short, concrete improvement tips.
"item_feedback" flags SPECIFIC weak bullets/lines (weak action verb, missing metric, vague, JD keyword
missing from this exact bullet) — up to 6 items, "section" must be one of: summary, experience, education,
skills, projects, certifications. "snippet" MUST be an exact short excerpt copied verbatim from that resume
so the UI can locate it — never paraphrase the snippet."""


async def score_document(content: ResumeDocumentContent, job_description: str = "") -> ResumeScoreResponse:
    resume_text = _flatten_content(content)
    user_prompt = f"RESUME:\n{resume_text[:6000]}"
    if job_description.strip():
        user_prompt += f"\n\nJOB DESCRIPTION:\n{job_description.strip()[:3000]}"
    data = await _openai_chat_json(_SCORE_SYSTEM_PROMPT, user_prompt, max_tokens=1200)
    item_feedback = []
    for item in (data.get("item_feedback") or [])[:6]:
        if not isinstance(item, dict):
            continue
        section = str(item.get("section", "")).strip().lower()
        snippet = str(item.get("snippet", "")).strip()
        issue = str(item.get("issue", "")).strip()
        if section and snippet and issue:
            item_feedback.append(ResumeScoreItemFeedback(section=section, snippet=snippet, issue=issue))
    return ResumeScoreResponse(
        score=max(0, min(100, int(data.get("score", 0)))),
        matched=[str(x) for x in (data.get("matched") or [])][:20],
        missing=[str(x) for x in (data.get("missing") or [])][:20],
        suggestions=[str(x) for x in (data.get("suggestions") or [])][:5],
        item_feedback=item_feedback,
    )


# ---------------------------------------------------------------------------
# Bullet writing (improve an existing bullet / turn free-text notes into one)
# Rules adapted from tailor/prompt/tool2_prompt.txt's BULLET RULES + tool3's
# scoring criteria (weak-verb / missing-metric deductions), so Maker bullets
# meet the same bar the Tailor pipeline enforces.
# ---------------------------------------------------------------------------

_BULLET_BASE_RULES = """You are an expert resume bullet writer. Write ONE resume bullet point.

RULES (matches the bar our resume-tailoring pipeline enforces):
- Follow the STAR method internally (Situation/Task -> Action -> Result), but output ONE tight sentence — never three.
- Start with a strong, specific action verb (e.g. "Architected", "Reduced", "Automated", "Shipped"). Never weak
  openers like "Responsible for", "Worked on", "Helped with", "Tasked with".
- Do NOT start with an action verb already used in these other bullets from the same section — pick a different one: {other_verbs}
- Include a quantified metric (%, $, time saved, scale, count) if the input plausibly supports one. Never invent
  a fake number — only quantify what's implied or stated.
- One sentence, ideally under 25 words. Readable and concrete — no jargon-stuffing, no buzzword soup.
- Preserve every concrete fact given (company names, technologies, numbers) verbatim — never invent achievements."""

_BULLET_WITH_JD = """
- A job description is provided below. Naturally weave in relevant keywords from it where truthful to the
  input — never fabricate experience the input doesn't support just to match a keyword.

JOB DESCRIPTION:
{job_description}"""

_BULLET_OUTPUT_FORMAT = """

Respond with JSON ONLY (no markdown fences): {{"bullet": "<the single rewritten bullet>"}}"""


def _bullet_system_prompt(other_bullets: list[str], job_description: str) -> str:
    other_verbs = ", ".join(b.split()[0] for b in other_bullets if b.strip()) or "(none yet)"
    prompt = _BULLET_BASE_RULES.format(other_verbs=other_verbs)
    if job_description.strip():
        prompt += _BULLET_WITH_JD.format(job_description=job_description.strip()[:3000])
    return prompt + _BULLET_OUTPUT_FORMAT


async def rewrite_bullet(
    bullet: str, context: str = "", other_bullets: list[str] | None = None, job_description: str = "",
) -> str:
    """Improve one existing bullet — same STAR/action-verb/metric bar as Tailor."""
    system = _bullet_system_prompt(other_bullets or [], job_description)
    user = f"Role/context: {context}\n\nCurrent bullet to improve:\n{bullet.strip()[:500]}"
    data = await _openai_chat_json(system, user, max_tokens=300)
    return str(data.get("bullet", "")).strip() or bullet


async def bullet_from_text(
    text: str, context: str = "", other_bullets: list[str] | None = None, job_description: str = "",
) -> str:
    """Turn a candidate's own raw, informal notes into a polished bullet under the same rules."""
    system = _bullet_system_prompt(other_bullets or [], job_description)
    user = f"Role/context: {context}\n\nCandidate's raw notes on what they did (turn this into one bullet):\n{text.strip()[:1000]}"
    data = await _openai_chat_json(system, user, max_tokens=300)
    return str(data.get("bullet", "")).strip()


# ---------------------------------------------------------------------------
# DOCX export (structure-aware — python-docx)
# ---------------------------------------------------------------------------

def generate_docx(content: ResumeDocumentContent, settings: ResumeDocumentSettings, out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    c = content.contact

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(c.full_name or "Untitled Resume")
    run.bold = True
    run.font.size = Pt(18)

    contact_line = " | ".join(x for x in [c.email, c.phone, c.location, c.linkedin, c.github, c.portfolio, c.website] if x)
    if contact_line:
        p = doc.add_paragraph(contact_line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def heading(text: str) -> None:
        h = doc.add_paragraph()
        r = h.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(12)

    if content.summary:
        heading("Summary")
        doc.add_paragraph(content.summary)

    section_map = {
        "experience": "Experience",
        "education": "Education",
        "skills": "Skills",
        "projects": "Projects",
        "certifications": "Certifications",
    }
    for key in content.section_order:
        label = section_map.get(key)
        if not label:
            continue
        if key == "experience" and content.experience:
            heading(label)
            for exp in content.experience:
                p = doc.add_paragraph()
                p.add_run(f"{exp.title} — {exp.company}").bold = True
                end = "Present" if exp.current else exp.end_date
                doc.add_paragraph(f"{exp.location}  |  {exp.start_date} – {end}")
                for b in exp.bullets:
                    doc.add_paragraph(b, style="List Bullet")
        elif key == "education" and content.education:
            heading(label)
            for edu in content.education:
                p = doc.add_paragraph()
                p.add_run(f"{edu.degree} {edu.field} — {edu.school}").bold = True
                doc.add_paragraph(f"{edu.location}  |  {edu.start_date} – {edu.end_date}" + (f"  |  GPA {edu.gpa}" if edu.gpa else ""))
        elif key == "skills" and content.skills:
            heading(label)
            for sk in content.skills:
                doc.add_paragraph(f"{sk.label}: {', '.join(sk.items)}")
        elif key == "projects" and content.projects:
            heading(label)
            for proj in content.projects:
                p = doc.add_paragraph()
                p.add_run(proj.name).bold = True
                if proj.description:
                    doc.add_paragraph(proj.description)
                for b in proj.bullets:
                    doc.add_paragraph(b, style="List Bullet")
        elif key == "certifications" and content.certifications:
            heading(label)
            for cert in content.certifications:
                doc.add_paragraph(f"{cert.name} — {cert.issuer} ({cert.date})")

    for custom in content.custom:
        heading(custom.title)
        for item in custom.items:
            doc.add_paragraph(item, style="List Bullet")

    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# PDF export (WeasyPrint — server-rendered HTML/CSS, flexbox-only for fidelity)
# ---------------------------------------------------------------------------

_ACCENT_COLORS = {
    "blue":    "#2563eb",
    "navy":    "#1e3a5f",
    "teal":    "#0d9488",
    "green":   "#059669",
    "emerald": "#10b981",
    "purple":  "#7c3aed",
    "orange":  "#ea580c",
    "amber":   "#d97706",
    "red":     "#dc2626",
    "slate":   "#475569",
}
_FONT_STACKS = {
    "serif": "'Georgia', 'Times New Roman', serif",
    "sans":  "'Helvetica Neue', Arial, sans-serif",
    "mono":  "'Courier New', monospace",
}
_ACCENT_TEMPLATES = {"modern", "modern-two-column", "vivid"}
_TWO_COLUMN_TEMPLATES = {"two-column", "modern-two-column", "vivid"}
_SIDEBAR_SECTIONS = {"education", "skills", "certifications"}
_MAIN_SECTIONS = {"experience", "projects"}
_SECTION_LABELS = {
    "experience": "Experience", "education": "Education", "skills": "Skills",
    "projects": "Projects", "certifications": "Certifications",
}


def _esc(text: str) -> str:
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _render_body_html(content: ResumeDocumentContent, settings: ResumeDocumentSettings) -> str:
    """The resume's inner HTML — shared by both the PDF export and the live
    screen preview, so the two can never visually drift apart. Page chrome
    (margins/size) is applied separately by each caller — see _wrap_pdf_html /
    _wrap_preview_html — since print (@page) and screen need different mechanisms."""
    c = content.contact
    accent = _ACCENT_COLORS.get(settings.accent_color, _ACCENT_COLORS["blue"])
    header_font = _FONT_STACKS.get(settings.header_font, _FONT_STACKS["sans"])
    body_font = _FONT_STACKS.get(settings.body_font, _FONT_STACKS["sans"])
    base_pt = 8 + settings.font_size_level          # 9-13pt body
    header_pt = base_pt + 4
    gap = 4 + settings.spacing_level * 2             # px between sections
    compact_mul = 0.75 if settings.compact else 1.0
    use_accent = settings.template in _ACCENT_TEMPLATES
    use_arrows = settings.template == "vivid"
    two_column = settings.template in _TWO_COLUMN_TEMPLATES
    heading_color = accent if use_accent else "#111827"
    heading_border = f"2px solid {accent}" if use_accent else \
        ("1px solid #9ca3af" if settings.template == "clean" else "1px solid #111827")
    heading_transform = "none" if settings.template == "latex" else "uppercase"

    contact_bits = [x for x in [c.email, c.phone, c.location, c.linkedin, c.github, c.portfolio, c.website] if x]
    contact_line = " &nbsp;|&nbsp; ".join(_esc(x) for x in contact_bits)

    def section_heading(label: str) -> str:
        return (
            f'<div style="font-family:{header_font};font-size:{base_pt+1}pt;font-weight:700;'
            f'color:{heading_color};text-transform:{heading_transform};letter-spacing:0.05em;'
            f'border-bottom:{heading_border};padding-bottom:2px;margin-top:{gap}px;'
            f'margin-bottom:{gap*compact_mul}px;">{_esc(label)}</div>'
        )

    def bullets_html(items: list[str]) -> str:
        if not items:
            return ""
        if use_arrows:
            rows = "".join(
                f'<div style="font-size:{base_pt}pt;display:flex;gap:6px;">'
                f'<span style="color:{accent};flex-shrink:0;">&#8594;</span><span>{_esc(b)}</span></div>'
                for b in items
            )
            return f'<div style="margin-top:2px;">{rows}</div>'
        items_html = "".join(f'<li style="font-size:{base_pt}pt;">{_esc(b)}</li>' for b in items)
        return f'<ul style="margin:2px 0 0 16px;padding:0;">{items_html}</ul>'

    def section_html(key: str) -> str:
        label = _SECTION_LABELS.get(key)
        if not label:
            return ""
        if key == "experience" and content.experience:
            body = "".join(
                f'<div style="margin-bottom:{gap*compact_mul}px;font-family:{body_font};">'
                f'<div style="display:flex;justify-content:space-between;font-weight:700;font-size:{base_pt}pt;">'
                f'<span>{_esc(exp.title)} — {_esc(exp.company)}</span>'
                f'<span>{_esc(exp.start_date)} – {_esc("Present" if exp.current else exp.end_date)}</span></div>'
                f'<div style="font-size:{base_pt-1}pt;color:#6b7280;">{_esc(exp.location)}</div>'
                f'{bullets_html(exp.bullets)}</div>'
                for exp in content.experience
            )
            return section_heading(label) + body
        if key == "education" and content.education:
            body = "".join(
                f'<div style="margin-bottom:{gap*compact_mul}px;font-family:{body_font};font-size:{base_pt}pt;">'
                f'<div style="font-weight:700;">{_esc(edu.degree)} {_esc(edu.field)}</div>'
                f'<div>{_esc(edu.school)}</div>'
                f'<div style="font-size:{base_pt-1}pt;color:#6b7280;">{_esc(edu.start_date)} – {_esc(edu.end_date)}'
                f'{f"  |  GPA {_esc(edu.gpa)}" if edu.gpa else ""}</div></div>'
                for edu in content.education
            )
            return section_heading(label) + body
        if key == "skills" and content.skills:
            body = "".join(
                f'<div style="font-family:{body_font};font-size:{base_pt}pt;margin-bottom:4px;">'
                f'<div style="font-weight:700;">{_esc(sk.label)}</div>'
                f'<div style="color:#374151;">{_esc(", ".join(sk.items))}</div></div>'
                for sk in content.skills
            )
            return section_heading(label) + body
        if key == "projects" and content.projects:
            body = "".join(
                f'<div style="margin-bottom:{gap*compact_mul}px;font-family:{body_font};font-size:{base_pt}pt;">'
                f'<b>{_esc(proj.name)}</b> {_esc(proj.description)}{bullets_html(proj.bullets)}</div>'
                for proj in content.projects
            )
            return section_heading(label) + body
        if key == "certifications" and content.certifications:
            body = "".join(
                f'<div style="font-family:{body_font};font-size:{base_pt}pt;margin-bottom:4px;">'
                f'{_esc(cert.name)} — {_esc(cert.issuer)} ({_esc(cert.date)})</div>'
                for cert in content.certifications
            )
            return section_heading(label) + body
        return ""

    header_html = (
        f'<div style="text-align:center;margin-bottom:{gap}px;">'
        f'<div style="font-family:{header_font};font-size:{header_pt}pt;font-weight:700;'
        f'color:{accent if settings.template in ("modern", "modern-two-column") else "#111827"};">{_esc(c.full_name)}</div>'
        f'<div style="font-family:{body_font};font-size:{base_pt-1}pt;color:#4b5563;">{_esc(c.headline)}</div>'
        f'<div style="font-family:{body_font};font-size:{base_pt-2}pt;color:#4b5563;margin-top:2px;">{contact_line}</div>'
        f'</div>'
    )

    summary_html = ""
    if content.summary:
        summary_html = section_heading("Summary") + (
            f'<div style="font-family:{body_font};font-size:{base_pt}pt;'
            f'margin-bottom:{gap if two_column else 0}px;">{_esc(content.summary)}</div>'
        )

    custom_html = "".join(
        section_heading(custom.title) + bullets_html(custom.items)
        for custom in content.custom
    )

    if not two_column:
        body_parts = [header_html, summary_html]
        body_parts += [section_html(key) for key in content.section_order]
        body_parts.append(custom_html)
        body_html = "\n".join(body_parts)
    else:
        sidebar_keys = [k for k in content.section_order if k in _SIDEBAR_SECTIONS]
        main_keys = [k for k in content.section_order if k in _MAIN_SECTIONS]
        sidebar_html = "".join(section_html(key) for key in sidebar_keys)
        main_html = "".join(section_html(key) for key in main_keys) + custom_html
        body_html = (
            header_html + summary_html +
            f'<div style="display:flex;gap:16px;">'
            f'<div style="flex:0 0 32%;">{sidebar_html}</div>'
            f'<div style="flex:1 1 68%;">{main_html}</div>'
            f'</div>'
        )

    return body_html


def _wrap_pdf_html(body_html: str, settings: ResumeDocumentSettings) -> str:
    """Full document for WeasyPrint — @page controls print margins/size."""
    page_size = "A4" if settings.page_size == "a4" else "Letter"
    return f"""<!doctype html>
<html><head><meta charset="utf-8"><style>
@page {{ size: {page_size}; margin: {settings.margin_top}mm {settings.margin_right}mm {settings.margin_bottom}mm {settings.margin_left}mm; }}
body {{ margin: 0; color: #111827; }}
ul {{ list-style: disc; }}
</style></head><body>{body_html}</body></html>"""


def _wrap_preview_html(body_html: str, settings: ResumeDocumentSettings) -> str:
    """Full document for an in-browser <iframe> preview. @page is print-only and
    invisible on screen, so page width/margins are applied here via ordinary CSS
    on a centered "paper" div instead — same visual result, different mechanism."""
    mm_to_px = lambda mm: round(mm / 25.4 * 96, 2)
    width_mm = 210 if settings.page_size == "a4" else 215.9
    width_px = mm_to_px(width_mm)
    pad = f"{mm_to_px(settings.margin_top)}px {mm_to_px(settings.margin_right)}px {mm_to_px(settings.margin_bottom)}px {mm_to_px(settings.margin_left)}px"
    return f"""<!doctype html>
<html><head><meta charset="utf-8"><style>
html, body {{ margin: 0; background: #e2e8f0; }}
body {{ display: flex; justify-content: center; padding: 16px 0; }}
.page {{ width: {width_px}px; min-height: {mm_to_px(279.4 if settings.page_size != 'a4' else 297)}px;
         background: #fff; box-shadow: 0 1px 4px rgba(0,0,0,0.15); padding: {pad}; box-sizing: border-box;
         color: #111827; }}
ul {{ list-style: disc; }}
</style></head><body><div class="page">{body_html}</div></body></html>"""


def generate_pdf(content: ResumeDocumentContent, settings: ResumeDocumentSettings, out_path: Path) -> Path:
    from weasyprint import HTML

    html = _wrap_pdf_html(_render_body_html(content, settings), settings)
    HTML(string=html).write_pdf(str(out_path))
    return out_path
